"""
Build a submission-ready .docx Results section with native Word tables and
embedded figures. Run from /projects/sandbox/BSRS.
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

# ---- base styles ----
normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(11)

def H(text, level=1):
    h = doc.add_heading(text, level=level)
    return h

def P(text, italic=False, size=11, color=None, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    r.bold = bold
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = RGBColor(*color)
    return p

def note(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    return p

def author_box(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(0x7a, 0x1f, 0x1f)
    return p

def table(headers, rows, caption=None, caption_above=True):
    if caption and caption_above:
        cp = doc.add_paragraph()
        cr = cp.add_run(caption)
        cr.bold = True
        cr.font.size = Pt(10)
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9.5)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(9.5)
    return t

def figure(path, width=6.2):
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

# ============================================================
H("4. Results", level=1)

author_box(
    "AUTHOR'S NOTE - READ BEFORE SUBMISSION (DELETE IN FINAL MANUSCRIPT). "
    "This section follows the JMIS reporting template and the SOP structural model "
    "(UE, UX exogenous; BSAT mediator; BSUC final endogenous; H1-H3 by PLS-SEM, P1 by "
    "fsQCA). (1) Measurement-model statistics in 4.1 come directly from the uploaded "
    "PLS-SEM analysis and are reusable. (2) The uploaded model actually used Attitude "
    "(ATT) as the outcome, not BSUC, and did NOT support H1-H3 (only a small NEGATIVE "
    "UE->ATT path, beta = -0.157; R2 ~ 0.03); the uploaded fsQCA returned "
    "~UE*~BSAT*(UX+BSUC) for ATT. Those actual results appear in Appendix A. (3) All "
    "values flagged [REPLACE] in 4.2-4.3 are ILLUSTRATIVE PLACEHOLDERS for the "
    "BSUC model and must be replaced with re-estimated SmartPLS / fsQCA output before "
    "submission. Do not publish placeholders as findings."
)

P("We report results in three stages. First, we evaluate the reflective measurement "
  "model (indicator reliability, internal consistency, convergent validity, and "
  "discriminant validity). Second, we assess the structural model and test the "
  "net-effect and mediation hypotheses H1-H3 using PLS-SEM. Third, we use fuzzy-set "
  "qualitative comparative analysis (fsQCA) to test proposition P1, which holds that "
  "multiple SDT-consistent configurations of UE, UX, and BSAT are jointly sufficient "
  "for high BSUC. The two methods are complementary: PLS-SEM estimates average net "
  "effects and the mediating role of brand satisfaction, whereas fsQCA examines "
  "configurational equifinality - whether several distinct recipes of the same "
  "antecedents can each produce brand success.")

P("The analysis draws on responses from 312 participants who each completed 21 "
  "reflective Likert items (1-7) measuring five latent constructs: User Engagement "
  "(UE; UE1-UE5), User Experience (UX; UX1-UX5), Brand Satisfaction (BSAT; "
  "BSAT1-BSAT4), Brand Success (BSUC; BSUC1-BSUC4), and Attitude (ATT; ATT_1-ATT_2). "
  "There were no missing data. The PLS-SEM algorithm was estimated with the "
  "path-weighting scheme; significance was obtained from 5,000 bootstrap subsamples. "
  "All indicator-block VIFs were below 3 (threshold = 5), indicating that "
  "collinearity and common-method concerns are not material.")

# ---- 4.1 ----
H("4.1 Measurement model", level=2)
P("Indicator reliability (outer loadings). All 21 indicators loaded on their "
  "theorized construct above the 0.708 benchmark (Hair et al., 2022). Loadings ranged "
  "from 0.88-0.91 (UE), 0.85-0.91 (UX), 0.83-0.88 (BSAT), 0.88-0.91 (BSUC), and "
  "~0.91 for both ATT items (Table 1).")

table(
    ["Construct", "Items", "Outer loading range", "All >= 0.708?"],
    [
        ["User Engagement (UE)", "UE1-UE5", "0.88 - 0.91", "Yes"],
        ["User Experience (UX)", "UX1-UX5", "0.85 - 0.91", "Yes"],
        ["Brand Satisfaction (BSAT)", "BSAT1-BSAT4", "0.83 - 0.88", "Yes"],
        ["Brand Success (BSUC)", "BSUC1-BSUC4", "0.88 - 0.91", "Yes"],
        ["Attitude (ATT)", "ATT_1, ATT_2", "~ 0.91", "Yes"],
    ],
    caption="Table 1. Indicator outer loadings by construct.",
)
note("Note. Loadings reported as the empirical range observed within each construct "
     "block. Replace with item-level loadings if required.")

P("Internal consistency and convergent validity. Cronbach's alpha ranges from 0.78 "
  "to 0.93 and composite reliability from 0.90 to 0.95 (both > 0.70); AVE ranges from "
  "0.73 to 0.82 (> 0.50). Reliability and convergent validity are established "
  "(Table 2; visualized in Figure 3).")

table(
    ["Construct", "Items", "Cronbach's a", "Composite rho_C", "AVE"],
    [
        ["User Engagement (UE)", "5", "0.93", "0.95", "0.79"],
        ["User Experience (UX)", "5", "0.92", "0.94", "0.77"],
        ["Brand Satisfaction (BSAT)", "4", "0.90", "0.95", "0.73"],
        ["Brand Success (BSUC)", "4", "0.90", "0.93", "0.77"],
        ["Attitude (ATT)", "2", "0.78", "0.90", "0.82"],
    ],
    caption="Table 2. Construct reliability and convergent validity.",
)
note("Note. All a and rho_C >= 0.70; all AVE >= 0.50 (Hair et al., 2022). Values "
     "taken directly from the uploaded PLS-SEM analysis.")

P("Discriminant validity. The square root of each construct's AVE (Table 3 diagonal) "
  "exceeds its correlations with other constructs (e.g., sqrt(AVE(UE)) = 0.89), and "
  "all HTMT ratios are below 0.85 (Table 4). Discriminant validity is supported.")

table(
    ["Construct", "UE", "UX", "BSAT", "BSUC", "ATT"],
    [
        ["UE", "0.889", "", "", "", ""],
        ["UX", "[REPLACE]", "0.877", "", "", ""],
        ["BSAT", "[REPLACE]", "[REPLACE]", "0.854", "", ""],
        ["BSUC", "[REPLACE]", "[REPLACE]", "[REPLACE]", "0.878", ""],
        ["ATT", "[REPLACE]", "[REPLACE]", "[REPLACE]", "[REPLACE]", "0.906"],
    ],
    caption="Table 3. Fornell-Larcker matrix (sqrt(AVE) on diagonal).",
)
note("Note. Diagonal = sqrt(AVE) from Table 2. Off-diagonal = inter-construct "
     "correlations (smaller than diagonal per source; insert from SmartPLS output).")

table(
    ["Construct", "UE", "UX", "BSAT", "BSUC", "ATT"],
    [
        ["UE", "-", "", "", "", ""],
        ["UX", "[REPLACE <0.85]", "-", "", "", ""],
        ["BSAT", "[REPLACE <0.85]", "[REPLACE <0.85]", "-", "", ""],
        ["BSUC", "[REPLACE <0.85]", "[REPLACE <0.85]", "[REPLACE <0.85]", "-", ""],
        ["ATT", "[REPLACE <0.85]", "[REPLACE <0.85]", "[REPLACE <0.85]", "[REPLACE <0.85]", "-"],
    ],
    caption="Table 4. Heterotrait-monotrait (HTMT) ratios.",
)
note("Note. Source confirms all HTMT < 0.85 but does not tabulate the matrix; insert values.")

P("Common method bias and collinearity. All indicator-block VIFs were below 3 "
  "(threshold = 5), so multicollinearity and common-method bias are unlikely to "
  "threaten the structural estimates. The measurement model satisfies all standard "
  "criteria, so we proceed to the structural model.")

# ---- 4.2 ----
H("4.2 Structural model and hypothesis testing", level=2)
author_box(
    "[REPLACE] All coefficients, t-values, p-values, R2, f2, and Q2 in this section "
    "are illustrative placeholders for the SOP (BSUC-outcome) model, which was not "
    "estimated in the uploaded analysis. Re-estimate (UE, UX -> BSAT -> BSUC; "
    "UE, UX -> BSUC) in SmartPLS and substitute actual values."
)

figure("figures/figure1_structural_model.png", width=6.2)
note("Figure 1. Structural model for avatar-mediated brand success. UE: User "
     "Engagement; UX: User Experience; BSAT: Brand Satisfaction; BSUC: Brand Success. "
     "Dashed lines = direct effects (H1, H2); solid lines = mediation chain (H3). "
     "Coefficients are [REPLACE] placeholders.")

P("Explanatory power (R2). The antecedents explain R2 = 0.46 [REPLACE] of the "
  "variance in BSAT and R2 = 0.55 [REPLACE] of the variance in BSUC (moderate to "
  "substantial; Hair et al., 2022).")
P("Paths to the mediator. UE -> BSAT (beta = 0.34 [REPLACE], t = 5.9, p < 0.001) and "
  "UX -> BSAT (beta = 0.41 [REPLACE], t = 7.2, p < 0.001) are both positive and "
  "significant.")
P("Direct effects (H1, H2). UE -> BSUC (beta = 0.23 [REPLACE], t = 3.9, p < 0.001) "
  "supports H1. UX -> BSUC (beta = 0.20 [REPLACE], t = 3.1, p < 0.01) supports H2. "
  "BSAT -> BSUC (beta = 0.39 [REPLACE], t = 6.6, p < 0.001) is the strongest single "
  "predictor.")

table(
    ["Path", "Std. beta", "t-value", "p-value", "Decision"],
    [
        ["UE -> BSAT", "0.34 [REPLACE]", "5.90", "< 0.001", "Significant"],
        ["UX -> BSAT", "0.41 [REPLACE]", "7.20", "< 0.001", "Significant"],
        ["UE -> BSUC (H1)", "0.23 [REPLACE]", "3.90", "< 0.001", "H1 supported"],
        ["UX -> BSUC (H2)", "0.20 [REPLACE]", "3.10", "0.002", "H2 supported"],
        ["BSAT -> BSUC", "0.39 [REPLACE]", "6.60", "< 0.001", "Significant"],
    ],
    caption="Table 5. Structural path coefficients (direct effects).",
)

P("Mediation (H3). Indirect effects are positive and significant: UE -> BSAT -> BSUC "
  "(beta = 0.13 [REPLACE], t = 4.1, p < 0.001) supports H3a; UX -> BSAT -> BSUC "
  "(beta = 0.16 [REPLACE], t = 4.8, p < 0.001) supports H3b. Because the direct UE "
  "and UX paths to BSUC remain significant with BSAT in the model, BSAT exhibits "
  "partial (complementary) mediation rather than full mediation (VAF ~ 37% and 44% "
  "[REPLACE], in the 20-80% partial range).")

table(
    ["Indirect path", "beta", "t-value", "p-value", "95% CI", "Mediation"],
    [
        ["UE -> BSAT -> BSUC (H3a)", "0.13 [REPLACE]", "4.10", "< 0.001", "[0.07, 0.20]", "Partial"],
        ["UX -> BSAT -> BSUC (H3b)", "0.16 [REPLACE]", "4.80", "< 0.001", "[0.09, 0.23]", "Partial"],
    ],
    caption="Table 6. Specific indirect (mediation) effects.",
)

P("Effect sizes (f2) and predictive relevance (Q2). BSAT has a medium-to-large effect "
  "on BSUC (f2 = 0.21 [REPLACE]); UE and UX have small-to-medium direct effects "
  "(f2 = 0.06 and 0.05 [REPLACE]). Blindfolding yielded Q2 > 0 for both endogenous "
  "constructs (Q2_BSAT = 0.31, Q2_BSUC = 0.38 [REPLACE]), confirming predictive "
  "relevance.")

table(
    ["Hypothesis", "Path", "Std. coeff.", "t-value", "p-value", "Supported?"],
    [
        ["H1", "UE -> BSUC (direct)", "0.23 [REPLACE]", "3.90", "< 0.001", "Yes"],
        ["H2", "UX -> BSUC (direct)", "0.20 [REPLACE]", "3.10", "0.002", "Yes"],
        ["H3a", "UE -> BSAT -> BSUC (indirect)", "0.13 [REPLACE]", "4.10", "< 0.001", "Yes"],
        ["H3b", "UX -> BSAT -> BSUC (indirect)", "0.16 [REPLACE]", "4.80", "< 0.001", "Yes"],
    ],
    caption="Table 7. Hypothesis testing summary.",
)
note("Note. Standardized coefficients; significance from 5,000 bootstrap subsamples. "
     "All [REPLACE] values pending re-estimation of the BSUC-outcome model.")

# ---- 4.3 ----
H("4.3 Configurational analysis (fsQCA)", level=2)
P("While PLS-SEM tests the net-effect hypotheses H1-H3, fsQCA was used to examine P1, "
  "which proposes that multiple SDT-consistent configurations of UE, UX, and BSAT are "
  "each sufficient for high BSUC. fsQCA models equifinality, conjunctural causation, "
  "and causal asymmetry, which net-effect regression cannot capture.")

P("Calibration. Each construct was averaged across items and calibrated into a fuzzy "
  "set using the direct method with anchors on the 1-7 scale: full out = 1 (0.00), "
  "crossover = 4 (0.50), full in = 7 (1.00) (Ragin, 2008).")

table(
    ["Variable", "Item(s)", "Full out (0.00)", "Crossover (0.50)", "Full in (1.00)"],
    [
        ["UE (Engagement)", "UE1-UE5 (avg)", "1", "4", "7"],
        ["UX (Experience)", "UX1-UX5 (avg)", "1", "4", "7"],
        ["BSAT (Satisfaction)", "BSAT1-BSAT4 (avg)", "1", "4", "7"],
        ["BSUC (outcome, SOP)", "BSUC1-BSUC4 (avg)", "1", "4", "7"],
    ],
    caption="Table 8. Variables and fuzzy-set calibration thresholds.",
)
note("Note. The uploaded analysis calibrated the outcome as Attitude (ATT). "
     "Recalibrate with BSUC as outcome to test P1.")

P("Necessity analysis. No single condition reached the 0.90 necessity threshold (all "
  "consistency <= 0.71; Table 9), so no individual antecedent is necessary for high "
  "BSUC - consistent with the configurational logic of P1.")

table(
    ["Condition", "Consistency", "Coverage"],
    [
        ["UE (Engagement)", "0.66", "0.63"],
        ["UX (Experience)", "0.69", "0.67"],
        ["BSAT (Satisfaction)", "0.68", "0.64"],
        ["BSUC*", "0.71", "0.65"],
    ],
    caption="Table 9. Analysis of necessary conditions.",
)
note("Note. No condition meets the 0.90 threshold. From the uploaded analysis "
     "(necessity vs. ATT); regenerate against BSUC.")

P("Sufficiency: truth table and solution. The fuzzy truth table was built over the "
  "logically possible configurations; rows were retained using a frequency threshold "
  "and a raw-consistency cutoff of 0.80, then minimized to the intermediate solution.")

table(
    ["UE", "UX", "BSAT", "BSUC", "N", "Raw consistency", "Raw coverage"],
    [
        ["0", "0", "0", "1", "25", "0.800", "0.120"],
        ["0", "1", "0", "0", "14", "0.786", "0.066"],
        ["...", "...", "...", "...", "...", "...", "..."],
    ],
    caption="Table 10. Truth-table rows observed in the data (uploaded analysis).",
)
note("Note. 1 = presence, 0 = absence. From the ATT-outcome analysis; regenerate for BSUC.")

author_box(
    "[REPLACE] The C1-C3 configurations below are illustrative placeholders consistent "
    "with P1. Replace with the actual intermediate-solution configurations from the "
    "BSUC-outcome fsQCA. The actual ATT-based solution is Table 12 / Appendix A."
)

figure("figures/figure2_fsqca_pathways.png", width=6.2)
note("Figure 2. Sufficient configurations for high brand success (fsQCA, P1): C1-C3 "
     "each lead to high BSUC, illustrating equifinality. Coverage/consistency are "
     "[REPLACE] placeholders.")

table(
    ["Config.", "UE", "UX", "BSAT", "Raw cov.", "Unique cov.", "Consistency"],
    [
        ["C1", "(*)", "(*)", "(*)", "0.46 [REPLACE]", "0.11", "0.89"],
        ["C2", "(*)", "o", "(*)", "0.41 [REPLACE]", "0.07", "0.87"],
        ["C3", "(*)", "(*)", "o", "0.38 [REPLACE]", "0.05", "0.85"],
        ["Solution", "", "", "", "0.61 [REPLACE]", "", "0.86"],
    ],
    caption="Table 11. Sufficient configurations for high brand success (P1).",
)
note("Note. (*) = presence of condition (high membership); o = condition not required "
     "(don't care). Solution coverage = 0.61, solution consistency = 0.86 [REPLACE]; "
     "each configuration exceeds the 0.80 benchmark. Visualized in Figure 2.")

P("Interpretation. C1 (UE, UX, BSAT all present): the all-strong recipe. C2 (UE and "
  "BSAT present): high engagement with high satisfaction suffices even when experience "
  "is not uniformly strong. C3 (UE and UX present): high engagement with high "
  "experience suffices even before satisfaction crystallizes. Several distinct "
  "configurations are each sufficient for high BSUC (solution consistency 0.86 > 0.80; "
  "solution coverage 0.61), supporting P1: brand success emerges from multiple "
  "SDT-consistent recipes rather than a single necessary driver.")

table(
    ["Solution term", "Raw consistency", "Raw coverage", "Unique coverage"],
    [
        ["~UE * ~BSAT * UX", "0.838", "0.424", "0.068"],
        ["~UE * ~BSAT * BSUC", "0.816", "0.448", "0.092"],
        ["Solution (UX + BSUC)", "0.813", "0.516", "-"],
    ],
    caption="Table 12. Actual solution terms reported in the uploaded fsQCA (ATT outcome).",
)
note("Note. Reproduced verbatim. This solution ~UE*~BSAT*(UX+BSUC) was estimated with "
     "ATT as outcome and CONTRADICTS the SOP/P1 framing; do not report it as evidence "
     "for P1. Use only to verify the re-estimation pipeline.")

# ---- Figure 3 ----
figure("figures/figure3_reliability.png", width=5.6)
note("Figure 3. Construct reliability and convergent validity (Cronbach's alpha, "
     "composite reliability rho_C, AVE) with 0.70 and 0.50 thresholds marked. Values "
     "from Table 2 (verified).")

# ---- Appendix ----
H("Appendix A. Actual estimates reported in the uploaded analyses (verification only)", level=2)
P("These are the empirical results the uploaded documents actually report. They model "
  "Attitude (ATT) as the outcome and do NOT support H1-H3 or P1 as framed in the SOP. "
  "They are retained so the author can reconcile the re-estimated BSUC model against "
  "the original pipeline.")

table(
    ["Path", "Coefficient", "t-value", "p-value"],
    [
        ["UE -> ATT", "-0.157", "2.79", "0.006"],
        ["UX -> ATT", "0.053", "0.93", "0.352"],
        ["BSAT -> ATT", "0.083", "1.47", "0.143"],
        ["BSUC -> ATT", "0.025", "0.45", "0.656"],
    ],
    caption="Table A1. Structural path coefficients as estimated (ATT outcome).",
)
note("Note. Only UE -> ATT significant, sign NEGATIVE. R2(ATT) ~ 0.03 (negligible); "
     "f2(UE->ATT) ~ 0.025 (small); Q2 <= 0 (no predictive relevance).")

P("Reporting conventions follow Hair et al. (2022) for PLS-SEM (loadings >= 0.708, "
  "a / rho_C >= 0.70, AVE >= 0.50, HTMT < 0.85, 5,000 bootstrap subsamples) and Ragin "
  "(2008) / Schneider & Wagemann (2012) for fsQCA (direct calibration; raw-consistency "
  "cutoff ~ 0.80; complex, parsimonious, and intermediate solutions with raw and "
  "unique coverage).", italic=True, size=9)

out = "Results_Section_JMIS.docx"
doc.save(out)
print("Saved", out, os.path.getsize(out), "bytes")
