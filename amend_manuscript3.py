#!/usr/bin/env python3
"""Third pass: align fsQCA captions 5/6 with the tables beneath them, fix two
in-text table references, and tidy remaining figure references."""
from docx import Document

SRC = "Manuscript Metaverse.docx"
doc = Document(SRC)


def find_p(sub):
    for p in doc.paragraphs:
        if sub in p.text:
            return p
    return None


def set_text(p, text):
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    p.add_run(text)


def rename(locate, new):
    p = find_p(locate)
    if p is None:
        print("  [MISS]", locate[:45]); return
    set_text(p, new)


def inline(locate, old, new):
    p = find_p(locate)
    if p is None or old not in p.text:
        print("  [MISS inline]", locate[:45]); return
    set_text(p, p.text.replace(old, new))


# Caption 5 sits above the subset/superset (Raw/Unique/consistency) table
rename("Table 5. Sufficiency (Intermediate Solution) Configurations for High B",
       "Table 5. Subset-Superset Analysis: Coverage and Consistency of Condition Combinations")
# Caption 6 sits above the intermediate-solution table (incl. single conditions)
rename("Table 6. Necessity (Subset/Superset) Analysis",
       "Table 6. Sufficiency Solution: Configurations and Single Conditions for High Brand Success")

# In-text reference fixes
inline("consistency values for each combination of variables are given in", "given in Table 7", "given in Table 5")
inline("presents the sufficient configurations for high brand success",
       "Table 7 presents the sufficient configurations",
       "Tables 6 and 7 jointly present the sufficient configurations")

doc.save(SRC)
print("Saved:", SRC)
