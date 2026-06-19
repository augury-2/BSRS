#!/usr/bin/env python3
"""Second amendment pass: fix the missed paragraph, standardize acronyms across
text + tables, reorder/renumber tables, remove the duplicate configuration
table, improve the configuration-table header, and insert the measurement-items
table (Table 1)."""
from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement

SRC = "Manuscript Metaverse.docx"
doc = Document(SRC)


def find_p(sub):
    for p in doc.paragraphs:
        if sub in p.text:
            return p
    return None


def find_all_p(sub):
    return [p for p in doc.paragraphs if sub in p.text]


def set_text(p, text):
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    p.add_run(text)
    return p


def inline(locate, old, new):
    p = find_p(locate)
    if p is None:
        print("  [MISS inline]", locate[:45]); return
    if old not in p.text:
        print("  [MISS old]", old[:45]); return
    set_text(p, p.text.replace(old, new))


def insert_after(p, text):
    new_p = OxmlElement('w:p')
    p._p.addnext(new_p)
    np = Paragraph(new_p, p._parent)
    if text:
        np.add_run(text)
    return np


# ---------- 1. Fix the missed "Sandbox" paragraph ----------
sb = None
for p in doc.paragraphs:
    if 'Sandbox' in p.text and 'voxel' in p.text:
        sb = p; break
if sb is not None:
    set_text(sb,
        "The Sandbox focuses on voxel-based content creation, enabling users to design, own, and "
        "monetise experiences, and has hosted brand activations such as Adidas NFT launches (Brandes "
        "and Dolp, 2025). Across these platforms, the study consistently classifies brand-relevant "
        "metaverse activity into four types: (1) digital-asset (NFT) activity, through which users "
        "acquire and display branded items linked to virtual identity (Lee and Shen, 2024); (2) social "
        "interaction, through which avatar-based community building occurs (Hennig-Thurau et al., "
        "2023); (3) content creation, through which users contribute to the virtual ecosystem "
        "(Bilgihan et al., 2024); and (4) virtual events. These activity types define the "
        "stratification used in sampling.")
    print("fixed Sandbox paragraph")


# ---------- 2. Inline number / grammar / figure fixes ----------
inline("presents the operationalization of these constructs",
       "Table 2 presents the operationalization", "Table 1 presents the operationalization")
inline("consolidated in Table 2 and Table 3",
       "consolidated in Table 2 and Table 3", "consolidated in Table 1 (measurement items and sources) and Table 3 (loadings and reliability)")
inline("Descriptive analysis was conducted using IBM SPSS", "(see Table 1)", "(see Table 2)")
inline("Demographic and Metaverse Engagement Details of Respondents", "Table 1:", "Table 2:")
inline("The presented Intermediate solution, as shown in Table", "as shown in Table 3", "as shown in Table 5")
inline("consistency values for each combination of variables are given in", "given in Table 4", "given in Table 7")
inline("subset-superset analysis was used to", "establish elucidate the link", "elucidate the link")
inline("yields five propositions, tested configurationally",
       "Figure 1 (conceptual model) summarises these expectations.",
       "The conceptual model (Figure 3) summarises these expectations.")


# ---------- 3. Delete duplicate UBS/UEX config table + orphan note ----------
dup = None
for t in doc.tables:
    cells_text = " ".join(c.text for r in t.rows for c in r.cells)
    if 'UEX' in cells_text or 'Triple Path' in cells_text:
        dup = t; break
if dup is not None:
    dup._element.getparent().remove(dup._element)
    print("deleted duplicate config table")

notes = find_all_p("Core condition present (essential for outcome)")
if len(notes) >= 2:
    notes[-1]._element.getparent().remove(notes[-1]._element)
    print("deleted orphan duplicate note")


# ---------- 4. Move reliability/loadings block before HTMT block ----------
htmt_cap = find_p("Heterotrait-Monotrait (HTMT) Discriminant Validity and Construct Association")
load_cap = find_p("Reliability and Convergent-Convergence Statistics")
load_note = find_p("UE = User Engagement; UX = User Experience; BSAT = Brand Satisfaction; BSUC = Brand Success. Source: Authors.")
load_tbl = None
for t in doc.tables:
    if any(c.text.strip() == 'Indicator' for r in t.rows for c in r.cells):
        load_tbl = t; break
if htmt_cap is not None and load_cap is not None and load_tbl is not None:
    htmt_cap._p.addprevious(load_cap._p)
    htmt_cap._p.addprevious(load_tbl._tbl)
    if load_note is not None:
        htmt_cap._p.addprevious(load_note._p)
    print("moved reliability block before HTMT block")
else:
    print("  [WARN] could not move reliability block", htmt_cap is not None, load_cap is not None, load_tbl is not None)


# ---------- 5. Global acronym standardization (text + tables) ----------
REPL = [('UMBS', 'BSAT'), ('MBSU', 'BSUC'), ('MVCX', 'UX'), ('UEX', 'UX'),
        ('UEN', 'UE'), ('UES', 'UE'), ('UBS', 'BSAT'), ('MBS', 'BSAT')]


def apply_acr(t):
    for a, b in REPL:
        t = t.replace(a, b)
    return t


cnt = 0
for p in doc.paragraphs:
    nt = apply_acr(p.text)
    if nt != p.text:
        set_text(p, nt); cnt += 1
for t in doc.tables:
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                nt = apply_acr(p.text)
                if nt != p.text:
                    set_text(p, nt)
print("acronym-normalized paragraphs:", cnt)


# ---------- 6. Improve configuration table header (Table 7) ----------
cfg = None
for t in doc.tables:
    if t.rows and t.rows[0].cells[0].text.strip() == 'Configuration' and len(t.columns) == 4:
        cfg = t; break
if cfg is not None:
    hdr = cfg.rows[0].cells
    labels = ['Configuration', 'C1: BSAT \u00b7 UX', 'C2: BSAT \u00b7 UE', 'C3: UX \u00b7 UE']
    for i, lab in enumerate(labels):
        set_text(hdr[i].paragraphs[0], lab)
        for rn in hdr[i].paragraphs[0].runs:
            rn.bold = True
    print("relabeled configuration table header")


# ---------- 7. Insert measurement-items table (Table 1) after Measures ----------
measures_p = find_p("User engagement was measured with ten items adapted from the User Engagement Scale")
if measures_p is not None:
    cap = insert_after(measures_p,
        "Table 1. Operationalization of Constructs and Measurement Items (seven-point Likert; "
        "BSAT = Brand Satisfaction; UE = User Engagement; UX = User Experience; BSUC = Brand Success)")
    items = [
        ["Construct", "Item", "Adapted wording", "Source"],
        ["BSAT", "BSAT1", "Overall, I am satisfied with this brand's presence in the metaverse.", "Fornell et al. (1996)"],
        ["BSAT", "BSAT2", "The brand's metaverse experience meets my expectations.", ""],
        ["BSAT", "BSAT3", "The brand's virtual offerings come close to my idea of an ideal brand experience.", ""],
        ["BSAT", "BSAT4", "I am satisfied with how the brand engages with me in virtual spaces.", ""],
        ["BSAT", "BSAT5", "My overall experience with the brand in the metaverse is satisfying.", ""],
        ["UE", "UE1", "I was absorbed in my activity within the metaverse. (focused attention)", "O'Brien & Toms (2010)"],
        ["UE", "UE2", "I lost track of time while engaging in the virtual environment. (focused attention)", ""],
        ["UE", "UE3", "I was so involved that I forgot about my immediate surroundings. (focused attention)", ""],
        ["UE", "UE4", "My interaction in the metaverse was fun. (felt involvement)", ""],
        ["UE", "UE5", "I felt involved in the brand's virtual activities. (felt involvement)", ""],
        ["UE", "UE6", "The experience was rewarding. (felt involvement)", ""],
        ["UE", "UE7", "The experience was worthwhile. (endurability)", ""],
        ["UE", "UE8", "I would recommend this metaverse activity to others. (endurability)", ""],
        ["UE", "UE9", "I was curious to explore the brand's virtual environment. (novelty)", ""],
        ["UE", "UE10", "The activities aroused my interest. (novelty)", ""],
        ["UX", "UX1", "I felt as though I was actually present in the virtual environment. (presence)", "Cadet & Chainay (2020); Basu et al. (2024)"],
        ["UX", "UX2", "The metaverse environment felt real to me. (presence)", ""],
        ["UX", "UX3", "I felt surrounded by the virtual world. (presence)", ""],
        ["UX", "UX4", "It was easy to move and navigate within the metaverse. (navigability)", ""],
        ["UX", "UX5", "Interacting with objects in the environment was intuitive. (navigability)", ""],
        ["UX", "UX6", "The environment responded the way I expected. (navigability)", ""],
        ["UX", "UX7", "The visual quality of the environment was high. (sensory quality)", ""],
        ["UX", "UX8", "The audio-visual richness enhanced my experience. (sensory quality)", ""],
        ["UX", "UX9", "The virtual environment authentically reflected real-world experience. (authenticity)", ""],
        ["UX", "UX10", "The brand's virtual space felt coherent and believable. (authenticity)", ""],
        ["BSUC", "BSUC1", "This brand stands out from competitors in the metaverse.", "Mathwick et al. (2008)"],
        ["BSUC", "BSUC2", "I am more likely to choose this brand because of its metaverse presence.", ""],
        ["BSUC", "BSUC3", "The brand has gained a competitive advantage through its metaverse activities.", ""],
        ["BSUC", "BSUC4", "I would advocate for this brand within virtual communities.", ""],
    ]
    tbl = doc.add_table(rows=len(items), cols=4)
    try:
        tbl.style = 'Light Grid Accent 1'
    except Exception:
        pass
    for ri, row in enumerate(items):
        for ci, val in enumerate(row):
            cell = tbl.rows[ri].cells[ci]
            set_text(cell.paragraphs[0], val)
            if ri == 0:
                for rn in cell.paragraphs[0].runs:
                    rn.bold = True
    cap._p.addnext(tbl._tbl)
    print("inserted measurement-items table (Table 1)")
else:
    print("  [WARN] measures paragraph not found for items table")


doc.save(SRC)
print("Saved:", SRC)
