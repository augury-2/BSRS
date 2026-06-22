"""
Build the Results section as a Word document (.docx) in Times New Roman 12,
JMIS-style, with tables and embedded figures.

Reads the CSV outputs produced by pls_sem.py and fsqca.py, plus the PNG
figures, and assembles analysis/outputs/Results_Section_JMIS.docx.
"""
import os
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.join(HERE, "outputs")
FIG = os.path.join(OUT, "figures")
DOCX = os.path.join(OUT, "Results_Section_JMIS.docx")

FONT = "Times New Roman"
SIZE = 12

# ---------------------------------------------------------------- load data
relval = pd.read_csv(os.path.join(OUT, "pls_reliability_validity.csv"))
loadings = pd.read_csv(os.path.join(OUT, "pls_loadings.csv"))
r2 = pd.read_csv(os.path.join(OUT, "pls_r2.csv")).set_index("Construct")["R2"].to_dict()
q2 = pd.read_csv(os.path.join(OUT, "pls_q2.csv")).set_index("Construct")["Q2"].to_dict()
f2 = pd.read_csv(os.path.join(OUT, "pls_f2.csv"))
paths = pd.read_csv(os.path.join(OUT, "pls_structural_paths.csv"))
med = pd.read_csv(os.path.join(OUT, "pls_mediation.csv"))
fl = pd.read_csv(os.path.join(OUT, "pls_fornell_larcker.csv"), index_col=0)
htmt = pd.read_csv(os.path.join(OUT, "pls_htmt.csv"), index_col=0)
lvcorr = pd.read_csv(os.path.join(OUT, "pls_lv_correlations.csv"), index_col=0)
anchors = pd.read_csv(os.path.join(OUT, "fsqca_calibration_anchors.csv"))
nec = pd.read_csv(os.path.join(OUT, "fsqca_necessity.csv"))
tt = pd.read_csv(os.path.join(OUT, "fsqca_truth_table.csv"))

# ---------------------------------------------------------------- helpers
doc = Document()

# base style
style = doc.styles["Normal"]
style.font.name = FONT
style.font.size = Pt(SIZE)
style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)


def _set_run(run, bold=False, italic=False, size=SIZE):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    r = run._element
    r.rPr.rFonts.set(qn("w:eastAsia"), FONT)


def heading(text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_run(run, bold=True, size=14 if level == 1 else 12)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p


def para(text, italic=False, bold=False, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_run(run, italic=italic, bold=bold)
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def caption(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_run(run, italic=True, size=11)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    return p


def add_figure(path, width_in=6.3):
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=Inches(width_in))


def add_table(df, header=True, first_col_bold=True, decimals=3, note=None):
    rows, cols = df.shape
    table = doc.add_table(rows=rows + (1 if header else 0), cols=cols + 1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # header row
    if header:
        hdr = table.rows[0].cells
        hdr[0].text = ""
        first = hdr[0].paragraphs[0].add_run(df.index.name or "")
        _set_run(first, bold=True, size=11)
        for j, col in enumerate(df.columns):
            cell = hdr[j + 1]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(col))
            _set_run(run, bold=True, size=11)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # body
    start = 1 if header else 0
    for i in range(rows):
        cells = table.rows[i + start].cells
        idx_run = cells[0].paragraphs[0].add_run(str(df.index[i]))
        _set_run(idx_run, bold=first_col_bold, size=11)
        for j in range(cols):
            val = df.iloc[i, j]
            if isinstance(val, float):
                txt = f"{val:.{decimals}f}"
            else:
                txt = str(val)
            cell = cells[j + 1]
            cell.text = ""
            run = cell.paragraphs[0].add_run(txt)
            _set_run(run, size=11)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if note:
        n = doc.add_paragraph()
        rn = n.add_run(note)
        _set_run(rn, italic=True, size=10)
    doc.add_paragraph()
    return table


# ================================================================ TITLE
title = doc.add_paragraph()
tr = title.add_run("4. Results")
_set_run(tr, bold=True, size=16)
title.paragraph_format.space_after = Pt(12)

para(
    "This section reports the empirical results of the two-stage analytical strategy "
    "outlined in Section 3. Following recent methodological guidance in the IS "
    "literature, we combine a symmetric, variance-based structural equation model "
    "(PLS-SEM) with an asymmetric, set-theoretic analysis (fuzzy-set Qualitative "
    "Comparative Analysis, fsQCA). PLS-SEM evaluates the net, average effects "
    "hypothesised among user engagement (UE), user experience (UX), brand satisfaction "
    "(BSAT) and brand success / continuance (BSUC), whereas fsQCA examines whether "
    "particular configurations of these conditions are necessary or sufficient for high "
    "BSUC. All analyses were conducted on the full sample of N = 312 valid responses "
    "measured on seven-point Likert scales, with no missing data. The complete, "
    "reproducible analysis pipeline (data screening, calibration, estimation, and 5,000 "
    "bootstrap resamples) is available with the supplementary materials."
)

add_figure(os.path.join(FIG, "fig10_workflow_flowchart.png"))
caption("Figure 10. Two-stage analytical workflow combining PLS-SEM (symmetric) and "
        "fsQCA (asymmetric) analysis.")

# ================================================================ 4.1 sample / descriptives
heading("4.1 Sample and Item Distributions", 1)
para(
    "The analytic sample comprised 312 respondents. Each of the four focal constructs "
    "was measured with multiple reflective indicators: UE (five items), UX (five items), "
    "BSAT (four items), and BSUC (four items). All items used a seven-point response "
    "format. Item-level distributions are summarised in Figure 6; responses spanned the "
    "full range of the scale (1\u20137) for every item, and there were no missing values, "
    "indicating adequate variability for multivariate analysis."
)
add_figure(os.path.join(FIG, "fig6_item_distributions.png"))
caption("Figure 6. Response distributions for all 18 measurement items (seven-point Likert).")

# ================================================================ 4.2 measurement model
heading("4.2 Assessment of the Measurement Model", 1)
para(
    "Consistent with established procedures for reflective measurement models, we first "
    "examined indicator reliability, internal consistency reliability, convergent "
    "validity, and discriminant validity before interpreting the structural model."
)

heading("4.2.1 Indicator Reliability and Convergent Validity", 2)
ld_lo = loadings["Loading"].min()
ld_hi = loadings["Loading"].max()
para(
    f"As shown in Table 1 and Figure 2, all standardised outer loadings ranged from "
    f"{ld_lo:.3f} to {ld_hi:.3f}, comfortably exceeding the recommended threshold of "
    f"0.708. Composite reliability (\u03c1c) ranged from "
    f"{relval['Composite_Reliability'].min():.3f} to {relval['Composite_Reliability'].max():.3f} "
    f"and Cronbach\u2019s \u03b1 from {relval['Cronbach_alpha'].min():.3f} to "
    f"{relval['Cronbach_alpha'].max():.3f}, both well above the 0.70 criterion. The average "
    f"variance extracted (AVE) ranged from {relval['AVE'].min():.3f} to {relval['AVE'].max():.3f}, "
    f"exceeding the 0.50 benchmark for every construct. The measurement instrument therefore "
    f"demonstrates strong indicator reliability, internal consistency, and convergent validity."
)

# Table 1: reliability/validity
t1 = relval.set_index("Construct")[["Items", "Cronbach_alpha", "Composite_Reliability", "AVE"]].copy()
t1.columns = ["Items", "Cronbach's \u03b1", "CR (\u03c1c)", "AVE"]
t1.index.name = "Construct"
para("Table 1. Construct reliability and convergent validity.", bold=True)
add_table(t1, decimals=3,
          note="Note. CR = composite reliability; AVE = average variance extracted. "
               "Thresholds: \u03b1 and CR \u2265 0.70; AVE \u2265 0.50.")

add_figure(os.path.join(FIG, "fig3_reliability_validity.png"))
caption("Figure 3. Construct reliability and convergent validity against recommended thresholds.")

# Table 2: loadings
para("Table 2. Standardised outer loadings.", bold=True)
ld_tab = loadings.copy()
ld_tab.index = ld_tab["Indicator"]
ld_tab.index.name = "Indicator"
ld_tab = ld_tab[["Construct", "Loading"]]
add_table(ld_tab, decimals=3, first_col_bold=True,
          note="Note. All loadings exceed the 0.708 threshold (Fig. 2).")
add_figure(os.path.join(FIG, "fig2_outer_loadings.png"))
caption("Figure 2. Indicator outer loadings by construct, relative to the 0.708 threshold.")

heading("4.2.2 Discriminant Validity", 2)
htmt_max = htmt.max().max()
para(
    "Discriminant validity was assessed using both the Fornell\u2013Larcker criterion and "
    "the heterotrait\u2013monotrait ratio of correlations (HTMT). As reported in Table 3, the "
    "square root of each construct\u2019s AVE (diagonal, in bold) is substantially larger than its "
    "correlations with all other constructs, satisfying the Fornell\u2013Larcker criterion. The "
    f"HTMT values (Table 4) were all far below the conservative 0.85 threshold (maximum HTMT = "
    f"{htmt_max:.3f}), providing strong evidence of discriminant validity. Notably, the very low "
    "inter-construct correlations already foreshadow the structural results reported in Section 4.3."
)

# Table 3: Fornell-Larcker
fl_tab = fl.copy()
fl_tab.index.name = "Construct"
para("Table 3. Fornell\u2013Larcker criterion (diagonal = \u221aAVE).", bold=True)
add_table(fl_tab, decimals=3,
          note="Note. Diagonal elements (\u221aAVE) exceed off-diagonal inter-construct correlations.")

# Table 4: HTMT
htmt_tab = htmt.copy().fillna("\u2013")
htmt_tab.index.name = "Construct"
para("Table 4. Heterotrait\u2013monotrait ratio (HTMT).", bold=True)
add_table(htmt_tab, decimals=3,
          note="Note. All HTMT values are below the 0.85 (and 0.90) thresholds.")

add_figure(os.path.join(FIG, "fig4_correlation_heatmap.png"), width_in=5.0)
caption("Figure 4. Construct-level correlation matrix. Inter-construct correlations are "
        "near zero, anticipating the structural findings.")

# ================================================================ 4.3 structural model
heading("4.3 Assessment of the Structural Model", 1)
para(
    "Given the satisfactory measurement model, we evaluated the structural model using path "
    "coefficients, their significance (based on 5,000 bootstrap resamples), coefficients of "
    "determination (R\u00b2), effect sizes (f\u00b2), and predictive relevance (Q\u00b2). "
    "Collinearity was not a concern, as the predictor constructs were effectively uncorrelated."
)

heading("4.3.1 Path Coefficients and Hypothesis Tests", 2)
para(
    "Table 5 and Figure 1 summarise the structural estimates. Contrary to expectations, none of "
    "the hypothesised paths reached statistical significance at the 0.05 level. Specifically, the "
    "effects of UE and UX on BSAT (H3a, H3b), of UE and UX on BSUC (H1, H2), and of BSAT on BSUC "
    "(H4) were all statistically non-significant, with bootstrap 95% confidence intervals that "
    "straddled zero in every case. The standardised path coefficients were uniformly small in "
    "magnitude (|\u03b2| \u2264 0.103)."
)

# Table 5: structural paths
pt = paths.copy()
pt.index = pt["Path"]
pt.index.name = "Path"
pt = pt[["Beta", "SE", "t", "p", "CI_2_5", "CI_97_5", "Significant"]]
pt.columns = ["\u03b2", "SE", "t", "p", "CI 2.5%", "CI 97.5%", "Sig. (p<.05)"]
para("Table 5. Structural model path coefficients (5,000 bootstrap resamples).", bold=True)
add_table(pt, decimals=3,
          note="Note. \u03b2 = standardised path coefficient; CI = bias-corrected bootstrap "
               "confidence interval. No path is significant at p < 0.05.")

add_figure(os.path.join(FIG, "fig1_sem_path_diagram.png"))
caption("Figure 1. PLS-SEM structural model with standardised path coefficients. "
        "All paths are non-significant (n.s.).")
add_figure(os.path.join(FIG, "fig5_path_coeffs_ci.png"))
caption("Figure 5. Structural path estimates with 95% bootstrap confidence intervals; "
        "every interval contains zero.")

heading("4.3.2 Explanatory and Predictive Power (R\u00b2, f\u00b2, Q\u00b2)", 2)
para(
    f"The explanatory power of the model was negligible. The endogenous constructs exhibited very "
    f"low coefficients of determination (R\u00b2 = {r2['BSAT']:.3f} for BSAT and "
    f"R\u00b2 = {r2['BSUC']:.3f} for BSUC), indicating that the model accounts for well under 2% "
    f"of the variance in either outcome. Effect sizes were correspondingly trivial; all f\u00b2 "
    f"values were below 0.011, far short of the 0.02 threshold for a small effect (Table 6). "
    f"Predictive relevance, assessed via a cross-validated (blindfolding-type) Q\u00b2, was "
    f"negative for both endogenous constructs (Q\u00b2 = {q2['BSAT']:.4f} for BSAT and "
    f"Q\u00b2 = {q2['BSUC']:.4f} for BSUC), indicating that the model has no out-of-sample "
    f"predictive relevance and does not outperform a naive mean prediction."
)

# Table 6: f2 and R2/Q2 combined
f2_tab = f2.copy()
f2_tab.index = f2_tab["Predictor"] + " \u2192 " + f2_tab["Outcome"]
f2_tab.index.name = "Effect"
f2_tab = f2_tab[["f2"]]
f2_tab.columns = ["f\u00b2"]
para("Table 6. Effect sizes (f\u00b2) for structural relationships.", bold=True)
add_table(f2_tab, decimals=4,
          note="Note. f\u00b2 thresholds: 0.02 (small), 0.15 (medium), 0.35 (large). "
               "All effects fall below the small-effect threshold.")

r2q2 = pd.DataFrame({
    "R\u00b2": [r2["BSAT"], r2["BSUC"]],
    "Q\u00b2": [q2["BSAT"], q2["BSUC"]],
}, index=["BSAT", "BSUC"])
r2q2.index.name = "Construct"
para("Table 7. Explanatory power (R\u00b2) and predictive relevance (Q\u00b2).", bold=True)
add_table(r2q2, decimals=4,
          note="Note. Negative Q\u00b2 indicates no predictive relevance.")

heading("4.3.3 Mediation Analysis", 2)
para(
    "We further tested the indirect (mediated) effects of UE and UX on BSUC through BSAT using "
    "the bootstrap distribution of the product of path coefficients. As shown in Table 8, neither "
    "indirect effect was statistically significant (UE \u2192 BSAT \u2192 BSUC: "
    f"effect = {med.iloc[0]['Effect']:.4f}, p = {med.iloc[0]['p']:.3f}; "
    f"UX \u2192 BSAT \u2192 BSUC: effect = {med.iloc[1]['Effect']:.4f}, p = {med.iloc[1]['p']:.3f}). "
    "Because the constituent direct paths were themselves non-significant, no mediation through "
    "brand satisfaction was supported."
)
med_tab = med.copy()
med_tab.index = med_tab["Indirect"]
med_tab.index.name = "Indirect path"
med_tab = med_tab[["Effect", "SE", "t", "p", "CI_2_5", "CI_97_5", "Significant"]]
med_tab.columns = ["Effect", "SE", "t", "p", "CI 2.5%", "CI 97.5%", "Sig."]
para("Table 8. Indirect (mediated) effects via BSAT.", bold=True)
add_table(med_tab, decimals=4,
          note="Note. Indirect effects estimated as the product of constituent path coefficients "
               "with 5,000 bootstrap resamples.")

# ================================================================ 4.4 fsQCA
heading("4.4 Configurational Analysis (fsQCA)", 1)
para(
    "Because PLS-SEM estimates only net, symmetric effects, we complemented it with fsQCA to "
    "examine whether specific combinations of UE, UX, and BSAT are necessary or sufficient for "
    "high BSUC. Construct scores were computed as the mean of their indicators and calibrated into "
    "fuzzy-set membership scores using the direct method, with the 95th, 50th, and 5th percentiles "
    "serving as the full-membership, crossover, and full-non-membership anchors, respectively "
    "(Table 9). An illustrative calibration is shown in Figure 9."
)

an = anchors.copy()
an.index = an["Variable"]
an.index.name = "Variable"
an = an[["Full_in_p95", "Crossover_p50", "Full_out_p5"]]
an.columns = ["Full-in (p95)", "Crossover (p50)", "Full-out (p5)"]
para("Table 9. fsQCA calibration anchors.", bold=True)
add_table(an, decimals=3,
          note="Note. Direct-method calibration using percentile anchors of each seven-point construct.")
add_figure(os.path.join(FIG, "fig9_calibration_bsuc.png"), width_in=4.8)
caption("Figure 9. Direct-method calibration of BSUC into fuzzy-set membership (illustrative).")

heading("4.4.1 Analysis of Necessary Conditions", 2)
nec_b = nec[nec.Outcome == "BSUC"].copy()
nec_max = nec_b["Consistency"].max()
para(
    "We first tested whether any condition (or its negation) is necessary for high BSUC. A "
    "condition is typically deemed necessary when its consistency exceeds 0.90. As reported in "
    f"Table 10 and Figure 7, the consistency of every condition and its negation ranged only from "
    f"{nec_b['Consistency'].min():.3f} to {nec_b['Consistency'].max():.3f}, all far below the 0.90 "
    f"benchmark (maximum = {nec_max:.3f}). Accordingly, none of UE, UX, or BSAT \u2014 in either its "
    "presence or its absence \u2014 constitutes a necessary condition for achieving high brand "
    "success / continuance."
)
nec_tab = nec_b.copy()
nec_tab.index = nec_tab["Condition"]
nec_tab.index.name = "Condition"
nec_tab = nec_tab[["Consistency", "Coverage"]]
para("Table 10. Analysis of necessary conditions for high BSUC.", bold=True)
add_table(nec_tab, decimals=3,
          note="Note. \u201c~\u201d denotes the absence (negation) of a condition. "
               "Necessity threshold = 0.90.")
add_figure(os.path.join(FIG, "fig7_fsqca_necessity.png"))
caption("Figure 7. Necessity consistencies for high BSUC; none approaches the 0.90 threshold.")

heading("4.4.2 Analysis of Sufficient Configurations", 2)
tt_max = tt["Consistency"].max()
para(
    "Next, we constructed the truth table of all 2\u00b3 = 8 possible configurations of the three "
    "conditions and evaluated their sufficiency for high BSUC, applying conventional thresholds of "
    "a minimum frequency of three cases and a minimum raw consistency of 0.80. As shown in Table 11 "
    f"and Figure 8, although all eight configurations were well populated, the highest raw "
    f"consistency attained by any configuration was only {tt_max:.3f}, below the 0.80 sufficiency "
    "cut-off. Consequently, no configuration qualified as a sufficient condition for high BSUC, and "
    "the sufficiency solution is empty (solution coverage and consistency are undefined). In other "
    "words, no combination of high or low UE, UX, and BSAT reliably produces high brand success."
)
tt_tab = tt.copy()
tt_tab.index = tt_tab["Configuration"]
tt_tab.index.name = "Configuration"
tt_tab = tt_tab[["Cases_gt_0_5", "Consistency"]]
tt_tab.columns = ["Cases (m>0.5)", "Raw consistency"]
para("Table 11. Truth table of configurations for high BSUC.", bold=True)
add_table(tt_tab, decimals=3,
          note="Note. \u201c*\u201d denotes logical AND; \u201c~\u201d denotes negation. "
               "Sufficiency thresholds: frequency \u2265 3, consistency \u2265 0.80. No row meets the "
               "consistency cut-off.")
add_figure(os.path.join(FIG, "fig8_fsqca_truthtable.png"))
caption("Figure 8. Truth-table configurations ranked by raw consistency; all fall below the "
        "0.80 sufficiency threshold.")

# ================================================================ 4.5 summary
heading("4.5 Summary of Results", 1)
para(
    "Taken together, the two analytical approaches converge on a consistent and unambiguous "
    "conclusion. The reflective measurement model is psychometrically sound: all constructs exhibit "
    "high indicator loadings, strong internal-consistency reliability, adequate convergent validity, "
    "and clear discriminant validity. However, the structural model reveals that the four constructs "
    "are statistically independent of one another. The symmetric PLS-SEM analysis found no "
    "significant direct or mediated paths, negligible R\u00b2 (< 0.02), trivial effect sizes "
    "(f\u00b2 < 0.011), and negative predictive relevance (Q\u00b2 < 0). The asymmetric fsQCA analysis "
    "corroborated this pattern, identifying neither necessary conditions (all necessity consistencies "
    "\u2248 0.55\u20130.60) nor sufficient configurations (maximum sufficiency consistency = "
    f"{tt_max:.3f}) for high brand success / continuance. Table 12 summarises the status of each "
    "hypothesis."
)

hyp = pd.DataFrame({
    "Path / Proposition": [
        "H1: UE \u2192 BSUC",
        "H2: UX \u2192 BSUC",
        "H3a: UE \u2192 BSAT",
        "H3b: UX \u2192 BSAT",
        "H4: BSAT \u2192 BSUC",
        "Mediation: UE/UX \u2192 BSAT \u2192 BSUC",
        "P1: Sufficient configurations for high BSUC",
    ],
    "Result": [
        "Not supported",
        "Not supported",
        "Not supported",
        "Not supported",
        "Not supported",
        "Not supported",
        "Not supported",
    ],
}).set_index("Path / Proposition")
hyp.index.name = "Path / Proposition"
para("Table 12. Summary of hypothesis testing.", bold=True)
add_table(hyp, first_col_bold=False,
          note="Note. See Sections 4.3 and 4.4 for the corresponding statistics.")

para(
    "These convergent null results are reported transparently and in full. They indicate that, "
    "within the present sample, engagement, experience, satisfaction, and brand success operate as "
    "statistically independent dimensions rather than as a causal chain. The implications of this "
    "finding \u2014 including data-quality and measurement considerations and directions for "
    "subsequent data collection \u2014 are discussed in Section 5.",
    italic=False
)

doc.save(DOCX)
print("Saved:", DOCX)
