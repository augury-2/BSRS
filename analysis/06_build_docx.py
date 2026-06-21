"""
06_build_docx.py
Render the Results section to a Word (.docx) research-paper document.
Body font: Times New Roman, 12 pt. Native Word tables; figures embedded at
publication size with numbered captions.
"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FIG = "analysis/figures"
OUT_DOCX = "Results_BSRS.docx"

doc = Document()

# ---------- global style: Times New Roman 12 ----------
def set_font(style_name, size, bold=False, color=None):
    st = doc.styles[style_name]
    st.font.name = "Times New Roman"
    st.font.size = Pt(size)
    st.font.bold = bold
    # ensure east-asian / complex scripts also use TNR
    rpr = st.element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rfonts.set(qn(attr), "Times New Roman")
    if color is not None:
        st.font.color.rgb = color

set_font("Normal", 12)
# headings -> TNR, black
for h, sz in [("Heading 1", 14), ("Heading 2", 13), ("Heading 3", 12)]:
    set_font(h, sz, bold=True, color=RGBColor(0, 0, 0))

normal = doc.styles["Normal"]
normal.paragraph_format.space_after = Pt(8)
normal.paragraph_format.line_spacing = 1.15


def add_para(text, italic=False, bold=False, align=None, size=12, space_after=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "justify":
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    # support simple **bold** inline markup
    import re
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            r = p.add_run(part[2:-2]); r.bold = True
        else:
            r = p.add_run(part)
        r.font.name = "Times New Roman"
        r.font.size = Pt(size)
        r.italic = italic
        if bold:
            r.bold = True
    return p


def add_heading(text, level):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = "Times New Roman"
        r.font.color.rgb = RGBColor(0, 0, 0)
    return h


def _set_cell_font(cell, bold=False, size=11, align=None):
    for p in cell.paragraphs:
        if align == "center":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(2)
        for r in p.runs:
            r.font.name = "Times New Roman"
            r.font.size = Pt(size)
            r.bold = bold


def _shade_cell(cell, hexcolor="D9D9D9"):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)


def add_table(rows, header=True, size=11, caption_above=None, note=None,
              col_align=None):
    if caption_above:
        cap = doc.add_paragraph()
        cap.paragraph_format.space_after = Pt(4)
        r = cap.add_run(caption_above)
        r.bold = True; r.font.name = "Times New Roman"; r.font.size = Pt(11)
    n_cols = max(len(r) for r in rows)
    t = doc.add_table(rows=0, cols=n_cols)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        cells = t.add_row().cells
        for j in range(n_cols):
            val = row[j] if j < len(row) else ""
            cells[j].text = str(val)
            is_header = header and i == 0
            al = "center"
            if col_align and j < len(col_align):
                al = col_align[j]
            _set_cell_font(cells[j], bold=is_header, size=size, align=al)
            if is_header:
                _shade_cell(cells[j], "DDDDDD")
    if note:
        n = doc.add_paragraph()
        n.paragraph_format.space_after = Pt(10)
        r = n.add_run(note)
        r.italic = True; r.font.name = "Times New Roman"; r.font.size = Pt(10)
    return t


def add_figure(path, caption, width_in=6.3):
    if not os.path.exists(path):
        add_para(f"[Figure not found: {path}]", italic=True)
        return
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run()
    run.add_picture(path, width=Inches(width_in))
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(12)
    r = cap.add_run(caption)
    r.font.name = "Times New Roman"; r.font.size = Pt(11); r.italic = True


# =====================================================================
# TITLE
# =====================================================================
title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Results")
r.bold = True; r.font.name = "Times New Roman"; r.font.size = Pt(16)
title.paragraph_format.space_after = Pt(12)

add_para(
    "This section reports the empirical findings in two stages. First, a "
    "variance-based structural equation model (partial least squares; PLS-SEM) is "
    "used to test the net, symmetric effects of the antecedents on brand-system "
    "continuance. Second, a fuzzy-set qualitative comparative analysis (fsQCA) is "
    "used to identify the configurations of conditions that are sufficient for high "
    "(and for low) continuance. The combination follows the established "
    "recommendation in the information-systems literature to triangulate "
    "correlational (PLS-SEM) and configurational (fsQCA) logics, because the two "
    "approaches answer complementary questions: PLS-SEM estimates the average "
    "marginal contribution of each antecedent, whereas fsQCA recovers the "
    "combinations of antecedents that are jointly sufficient for the outcome and "
    "explicitly allows for equifinality and causal asymmetry (Ragin, 2008; "
    "Woodside, 2013; Pappas & Woodside, 2021).", align="justify")

add_para("The constructs are defined as follows.")
add_table(
    [["Code", "Construct", "Items"],
     ["UE", "User Engagement", "UE1\u2013UE5"],
     ["UX", "User Experience", "UX1\u2013UX5"],
     ["BSAT", "Brand-System Satisfaction", "BSAT1\u2013BSAT4"],
     ["BSUC", "Brand-System Use Continuance (outcome)", "BSUC1\u2013BSUC4"]],
    col_align=["center", "left", "center"])

add_para(
    "Note on the two ATT items. The instrument also contained two items labelled "
    "ATT_1 and ATT_2. Inspection of their distributions and inter-item behaviour "
    "shows that they do not behave as a reflective construct: ATT_1 is almost "
    "invariant (M = 5.90, SD = 0.53), ATT_2 is low and right-skewed (M = 2.10), and "
    "the two correlate at r = \u2212.005 with each other and \u2248 0 with every substantive "
    "item (Figure A1). This is the signature of attention / quality-control items "
    "rather than a latent variable. They were therefore used only for data "
    "screening and were excluded from the measurement and structural models and "
    "from the fsQCA. A robustness model that forced them into a latent variable "
    "produced a construct with no convergent validity and changed none of the "
    "substantive conclusions.", italic=True, align="justify")

# =====================================================================
# 1. SCREENING
# =====================================================================
add_heading("1. Sample, data screening, and common method bias", 1)
add_para(
    "The analysis is based on N = 312 complete responses measured on seven-point "
    "Likert scales. There were no missing values across the 20 items and no "
    "straight-line response patterns, so the full sample was retained. Because "
    "PLS-SEM is a non-parametric method, multivariate normality is not required; "
    "nonetheless, Shapiro\u2013Wilk tests rejected univariate normality for all items "
    "(p < .001), confirming that a variance-based estimator is appropriate over "
    "covariance-based SEM (Hair et al., 2019). Item-level descriptive statistics "
    "are reported in Table 1.", align="justify")

add_table(
    [["Item", "Mean", "SD", "Median", "Min", "Max", "Skewness", "Kurtosis"]] +
    [[it, "4.00", "2.01", "4", "1", "7", "0.00", "\u22121.25"]
     for it in ["UE1","UE2","UE3","UE4","UE5","UX1","UX2","UX3","UX4","UX5",
                "BSAT1","BSAT2","BSAT3","BSAT4","BSUC1","BSUC2","BSUC3","BSUC4"]],
    caption_above="Table 1. Item descriptive statistics (N = 312).", size=10,
    note="Full numeric precision is provided in analysis/outputs/T1_item_descriptives.csv.")

add_para(
    "The sampling adequacy of the item set was excellent (Kaiser\u2013Meyer\u2013Olkin = "
    "0.903) and Bartlett's test of sphericity was significant (\u03c7\u00b2(153) = 3300.5, "
    "p < .001). An exploratory factor analysis (principal-axis extraction, promax "
    "rotation) recovered a clean four-factor structure mapping exactly onto UE, UX, "
    "BSAT and BSUC, with every substantive item loading > 0.77 on its intended "
    "factor and < 0.13 on all others.", align="justify")
add_para(
    "Common method bias. Harman's single-factor test indicated that the first "
    "unrotated factor accounted for 34.4% of the variance\u2014below the 50% "
    "threshold\u2014so common method variance is unlikely to be a serious threat "
    "(Podsakoff et al., 2003). This is corroborated by the full collinearity "
    "assessment in the structural model (all inner VIF \u2264 1.35; Kock, 2015) and by "
    "the discriminant-validity evidence below.", align="justify")
add_para(
    "Transparency note on data provenance. The 18 substantive items each exhibit "
    "identical marginal moments (mean = 4.000, SD = 2.006, skew = 0, excess "
    "kurtosis = \u22121.25), which are exactly the moments of a discrete uniform "
    "distribution on {1,\u2026,7}. Combined with a realistic, block-structured "
    "correlation matrix (Figure A1), this is consistent with data generated by a "
    "Gaussian-copula-type simulation rather than raw field data. The analyses are "
    "reported exactly as the data dictate; if these are synthetic or pilot data, "
    "the substantive interpretation should be treated as illustrative of the "
    "analytical pipeline rather than as confirmed empirical fact.",
    italic=True, align="justify")

# =====================================================================
# 2. MEASUREMENT MODEL
# =====================================================================
add_heading("2. Measurement model assessment", 1)
add_para(
    "Reflective measurement quality was assessed following the standard sequence "
    "for PLS-SEM (Hair et al., 2019, 2022): indicator reliability, internal "
    "consistency reliability, convergent validity, and discriminant validity. "
    "Estimation used Mode A (reflective) measurement with the path weighting "
    "scheme. As an external check, the measurement model was re-estimated with the "
    "open-source plspm package; the two implementations agreed to within 0.003 on "
    "every loading and produced identical reliability and AVE values.",
    align="justify")

add_heading("2.1 Indicator reliability and convergent validity", 2)
add_para(
    "All standardized outer loadings lie between 0.812 and 0.879 (Figure 7), above "
    "the 0.708 benchmark. Average variance extracted (AVE) ranges from 0.692 to "
    "0.747, exceeding 0.50 and establishing convergent validity (Table 2). Internal "
    "consistency is strong but not excessive: Cronbach's \u03b1 (0.885\u20130.892), the "
    "Dijkstra\u2013Henseler \u03c1_A (0.887\u20130.895), and composite reliability \u03c1_c "
    "(0.918\u20130.922) all fall in the recommended 0.70\u20130.95 band. Indicator-level "
    "collinearity is unproblematic (maximum outer VIF = 2.46).", align="justify")

add_table(
    [["Construct", "# items", "Cronbach's \u03b1", "\u03c1_A", "CR (\u03c1_c)", "AVE"],
     ["UE", "5", "0.889", "0.892", "0.918", "0.692"],
     ["UX", "5", "0.892", "0.895", "0.921", "0.699"],
     ["BSAT", "4", "0.885", "0.887", "0.921", "0.743"],
     ["BSUC", "4", "0.887", "0.888", "0.922", "0.747"]],
    caption_above="Table 2. Reliability and convergent validity.")
add_figure(f"{FIG}/Fig7_outer_loadings.png",
           "Figure 7. Indicator outer loadings by construct (dashed line = 0.708 threshold).")

add_heading("2.2 Discriminant validity", 2)
add_para(
    "Discriminant validity was established with three complementary criteria. "
    "Fornell\u2013Larcker (Table 3, Figure 6): for every construct the square root of "
    "the AVE (diagonal) exceeds its correlations with all other constructs. "
    "Heterotrait\u2013monotrait ratio (HTMT; Table 4, Figure 5): all values fall between "
    "0.30 and 0.615, far below the conservative 0.85 threshold (Henseler et al., "
    "2015), with bootstrap confidence intervals excluding 1. Cross-loadings: every "
    "indicator loads highest on its own construct by a margin of at least 0.40.",
    align="justify")

add_table(
    [["", "UE", "UX", "BSAT", "BSUC"],
     ["UE", "0.832", "", "", ""],
     ["UX", "0.365", "0.836", "", ""],
     ["BSAT", "0.300", "0.450", "0.862", ""],
     ["BSUC", "0.265", "0.341", "0.545", "0.865"]],
    caption_above="Table 3. Fornell\u2013Larcker matrix (diagonal = \u221aAVE).",
    col_align=["left", "center", "center", "center", "center"])

add_table(
    [["", "UE", "UX", "BSAT", "BSUC"],
     ["UE", "\u2014", "", "", ""],
     ["UX", "0.407", "\u2014", "", ""],
     ["BSAT", "0.334", "0.503", "\u2014", ""],
     ["BSUC", "0.298", "0.381", "0.615", "\u2014"]],
    caption_above="Table 4. Heterotrait\u2013monotrait (HTMT) ratios.",
    col_align=["left", "center", "center", "center", "center"])

add_figure(f"{FIG}/Fig5_HTMT.png", "Figure 5. Heterotrait\u2013monotrait (HTMT) ratios.", 4.8)
add_figure(f"{FIG}/Fig6_FornellLarcker.png",
           "Figure 6. Fornell\u2013Larcker matrix (diagonal = \u221aAVE; off-diagonal = construct correlations).", 4.8)

# =====================================================================
# 3. STRUCTURAL MODEL
# =====================================================================
add_heading("3. Structural model assessment", 1)
add_heading("3.1 Collinearity and model fit", 2)
add_para(
    "Inner-model collinearity is negligible: all predictor-side VIFs lie between "
    "1.15 and 1.35 (Table 5), ruling out lateral collinearity and supporting the "
    "conclusion that common method bias does not distort the path estimates "
    "(Kock, 2015). The standardized root mean square residual (SRMR = 0.047) is "
    "below the 0.08 cutoff, indicating good approximate fit.", align="justify")
add_table(
    [["Predictor \u2192 Outcome", "VIF"],
     ["UE \u2192 BSAT", "1.15"], ["UX \u2192 BSAT", "1.15"],
     ["UE \u2192 BSUC", "1.19"], ["UX \u2192 BSUC", "1.35"],
     ["BSAT \u2192 BSUC", "1.29"]],
    caption_above="Table 5. Inner-model collinearity (VIF).",
    col_align=["left", "center"])

add_heading("3.2 Explanatory and predictive power", 2)
add_para(
    "The model explains a meaningful share of variance (Table 6, Figure 3): "
    "R\u00b2 = 0.224 for satisfaction (BSAT) and R\u00b2 = 0.315 for continuance (BSUC). "
    "Blindfolding (omission distance = 7) yielded cross-validated redundancy "
    "Q\u00b2 = 0.160 (BSAT) and 0.228 (BSUC), both above zero, establishing predictive "
    "relevance.", align="justify")
add_table(
    [["Construct", "R\u00b2", "Adjusted R\u00b2", "Q\u00b2 (blindfolding)"],
     ["BSAT", "0.224", "0.219", "0.160"],
     ["BSUC", "0.315", "0.308", "0.228"]],
    caption_above="Table 6. Variance explained and predictive relevance.")
add_figure(f"{FIG}/Fig3_R2_Q2.png",
           "Figure 3. Explanatory (R\u00b2) and predictive (Q\u00b2) power of the endogenous constructs.", 5.0)

add_para(
    "A formal out-of-sample PLSpredict analysis (Shmueli et al., 2019; 10-fold "
    "cross-validation) confirmed predictive validity for the key target construct: "
    "Q\u00b2_predict was positive for all four BSUC indicators (0.035\u20130.067), and the "
    "PLS path model produced a lower RMSE than the naive linear-model benchmark for "
    "three of the four indicators (Table 7), indicating at least medium "
    "out-of-sample predictive power.", align="justify")
add_table(
    [["Indicator", "Q\u00b2_predict", "RMSE (PLS)", "RMSE (LM)", "PLS < LM?"],
     ["BSUC1", "0.052", "1.951", "1.946", "No"],
     ["BSUC2", "0.057", "1.945", "1.981", "Yes"],
     ["BSUC3", "0.036", "1.967", "2.006", "Yes"],
     ["BSUC4", "0.067", "1.935", "1.953", "Yes"]],
    caption_above="Table 7. PLSpredict (10-fold CV) for the BSUC indicators.")

add_heading("3.3 Path coefficients and hypothesis tests", 2)
add_para(
    "Path significance was assessed with bootstrapping (5,000 resamples); Table 8 "
    "reports standardized coefficients, standard errors, t-statistics, two-tailed "
    "p-values, 95% percentile confidence intervals, and f\u00b2 effect sizes. The "
    "estimated model is depicted in Figure 1 and the path estimates with confidence "
    "intervals in Figure 2.", align="justify")
add_table(
    [["H", "Path", "\u03b2", "SE", "t", "p", "95% CI", "f\u00b2", "Decision"],
     ["H1", "UE \u2192 BSAT", "0.157", "0.054", "2.89", ".004", "[0.052, 0.266]", "0.028", "Supported"],
     ["H2", "UX \u2192 BSAT", "0.393", "0.052", "7.59", "<.001", "[0.291, 0.494]", "0.172", "Supported"],
     ["H5", "BSAT \u2192 BSUC", "0.476", "0.047", "10.15", "<.001", "[0.382, 0.566]", "0.256", "Supported"],
     ["H3", "UE \u2192 BSUC", "0.088", "0.048", "1.82", ".069", "[\u22120.004, 0.185]", "0.010", "Not supported"],
     ["H4", "UX \u2192 BSUC", "0.095", "0.054", "1.78", ".076", "[\u22120.008, 0.203]", "0.010", "Not supported"]],
    caption_above="Table 8. Structural path coefficients (bootstrap, 5,000 resamples).",
    size=10, col_align=["center","left","center","center","center","center","center","center","center"])
add_para(
    "Three results stand out. First, satisfaction is the proximal driver of "
    "continuance: BSAT \u2192 BSUC is the strongest path (\u03b2 = 0.476, p < .001) with a "
    "medium-to-large effect (f\u00b2 = 0.256). Second, experience dominates engagement "
    "as an antecedent of satisfaction: UX \u2192 BSAT (\u03b2 = 0.393) is roughly 2.5\u00d7 the "
    "size of UE \u2192 BSAT (\u03b2 = 0.157). Third, neither antecedent exerts a significant "
    "direct effect on continuance once satisfaction is accounted for (H3, H4), "
    "foreshadowing full mediation.", align="justify")
add_figure(f"{FIG}/Fig1_structural_model.png",
           "Figure 1. PLS-SEM structural model with standardized path coefficients "
           "(solid = significant at p < .05; dashed = n.s.; *p < .05, **p < .01, ***p < .001).")
add_figure(f"{FIG}/Fig2_path_CIs.png",
           "Figure 2. Path estimates with 95% bootstrap confidence intervals (5,000 resamples).", 5.5)
add_figure(f"{FIG}/Fig4_f2.png",
           "Figure 4. f\u00b2 effect sizes for all structural paths (reference lines at 0.02, 0.15, 0.35).", 5.2)

add_heading("3.4 Mediation analysis", 2)
add_para(
    "Following Zhao et al. (2010) and Nitzl et al. (2016), mediation was tested on "
    "the bootstrapped indirect effects. Both specific indirect paths through "
    "satisfaction are significant (Table 9): UE \u2192 BSAT \u2192 BSUC (indirect = 0.075, "
    "p = .007) and UX \u2192 BSAT \u2192 BSUC (indirect = 0.187, p < .001). Because in each "
    "case the indirect effect is significant while the direct effect is not, the "
    "data indicate full (indirect-only) mediation: satisfaction fully transmits the "
    "influence of both experience and engagement onto continuance.", align="justify")
add_table(
    [["Indirect path", "Indirect", "SE", "t", "p", "95% CI", "Direct", "Total", "Type"],
     ["UE \u2192 BSAT \u2192 BSUC", "0.075", "0.027", "2.72", ".007", "[0.024, 0.132]", "0.088 (n.s.)", "0.162", "Full"],
     ["UX \u2192 BSAT \u2192 BSUC", "0.187", "0.030", "6.15", "<.001", "[0.131, 0.251]", "0.095 (n.s.)", "0.282", "Full"]],
    caption_above="Table 9. Specific indirect effects and total effects on BSUC.",
    size=10, col_align=["left","center","center","center","center","center","center","center","center"])
add_para(
    "The total effects on continuance rank the antecedents as BSAT (0.476) > "
    "UX (0.282) > UE (0.162), reinforcing that satisfaction is the dominant lever "
    "and experience the dominant upstream antecedent.", align="justify")

add_heading("3.5 Importance\u2013performance map analysis (IPMA)", 2)
add_para(
    "An IPMA was conducted with BSUC as the target. The importance dimension (total "
    "effects) reproduces the ranking BSAT > UX > UE (Figure 8). The performance "
    "dimension is uninformative in these data because every construct's rescaled "
    "mean sits at the scale midpoint (\u2248 50/100), a direct consequence of the "
    "uniform item marginals; the IPMA is therefore reported for completeness, but "
    "performance-based prioritization should not be inferred from this sample.",
    align="justify")
add_figure(f"{FIG}/Fig8_IPMA.png",
           "Figure 8. Importance\u2013performance map analysis (target: BSUC).", 5.2)

# =====================================================================
# 4. fsQCA
# =====================================================================
add_heading("4. Configurational analysis (fsQCA)", 1)
add_para(
    "To complement the variance-based analysis with a configurational and "
    "asymmetric view\u2014which combinations of conditions are sufficient for high "
    "continuance, and whether the recipe for low continuance is simply the mirror "
    "image\u2014an fsQCA was conducted with UE, UX and BSAT as conditions and BSUC as "
    "the outcome (Ragin, 2008; Fiss, 2011; Pappas & Woodside, 2021).", align="justify")

add_heading("4.1 Calibration", 2)
add_para(
    "Construct scores (item means) were calibrated into fuzzy-set membership using "
    "the direct method, with anchors set at the 95th percentile (full membership), "
    "the 50th percentile (cross-over), and the 5th percentile (full non-membership) "
    "of each construct (Table 10). Cases on the cross-over were nudged by 0.001 to "
    "avoid loss (Ragin, 2008). Figure 9 overlays the raw distributions with the "
    "membership functions.", align="justify")
add_table(
    [["Construct", "Full-in (P95)", "Cross-over (P50)", "Full-out (P5)"],
     ["UE", "6.60", "3.90", "1.40"],
     ["UX", "6.80", "4.00", "1.20"],
     ["BSAT", "6.61", "4.25", "1.25"],
     ["BSUC", "6.75", "4.00", "1.25"]],
    caption_above="Table 10. Calibration anchors (direct method, percentile-based).")
add_figure(f"{FIG}/Fig9_calibration.png",
           "Figure 9. fsQCA calibration: raw composite distributions (bars) and fuzzy-set membership functions (curves).")

add_heading("4.2 Analysis of necessary conditions", 2)
add_para(
    "No single condition is necessary for high continuance: the highest necessity "
    "consistency is for BSAT (0.751), below the 0.90 threshold (Table 11, "
    "Figure 11). The same holds for the negated outcome (highest: ~BSAT, 0.784). "
    "The absence of any necessary condition is itself informative\u2014high continuance "
    "does not require any one antecedent in isolation, motivating the sufficiency "
    "analysis.", align="justify")
add_table(
    [["Condition", "Cons. (high BSUC)", "Cov. (high BSUC)", "Cons. (~BSUC)", "Cov. (~BSUC)"],
     ["UE", "0.698", "0.685", "0.576", "0.567"],
     ["~UE", "0.559", "0.568", "0.680", "0.693"],
     ["UX", "0.718", "0.714", "0.560", "0.559"],
     ["~UX", "0.557", "0.558", "0.713", "0.717"],
     ["BSAT", "0.751", "0.776", "0.476", "0.494"],
     ["~BSAT", "0.511", "0.493", "0.784", "0.759"]],
    caption_above="Table 11. Analysis of necessary conditions (consistency threshold = 0.90).",
    size=10, col_align=["left","center","center","center","center"])
add_figure(f"{FIG}/Fig11_fsqca_necessity.png",
           "Figure 11. fsQCA necessity XY plots (no condition reaches the 0.90 necessity threshold).")

add_heading("4.3 Analysis of sufficient configurations", 2)
add_para(
    "A truth table was constructed over the 2\u00b3 = 8 corners of the vector space. "
    "Rows were coded as sufficient using a frequency threshold of 4 (all rows "
    "exceeded this; total retained cases = 312), a raw-consistency threshold of "
    "0.80, and a PRI threshold of 0.70 (Greckhamer et al., 2018). The truth table "
    "is summarized in Table 12 and Figure 12.", align="justify")
add_table(
    [["UE", "UX", "BSAT", "n cases", "Raw cons.", "PRI", "Sufficient?"],
     ["1", "0", "1", "19", "0.883", "0.698", "No (PRI < 0.70)"],
     ["1", "1", "1", "74", "0.857", "0.737", "Yes"],
     ["0", "1", "1", "34", "0.853", "0.664", "No"],
     ["0", "0", "1", "30", "0.817", "0.542", "No"],
     ["1", "1", "0", "27", "0.764", "0.443", "No"],
     ["0", "1", "0", "26", "0.744", "0.366", "No"],
     ["1", "0", "0", "36", "0.687", "0.303", "No"],
     ["0", "0", "0", "66", "0.570", "0.190", "No"]],
    caption_above="Table 12. Truth table (sorted by raw consistency).", size=10)
add_para(
    "The Quine\u2013McCluskey minimisation returns a single sufficient configuration; "
    "the complex, parsimonious and intermediate solutions coincide (Table 13). All "
    "three conditions appear and\u2014because retained in both the parsimonious and "
    "intermediate solution\u2014each is a core condition (Fiss, 2011). High brand-system "
    "continuance arises when engagement, experience and satisfaction are all "
    "simultaneously high (solution coverage = 0.497, consistency = 0.857). The "
    "corner UE\u00b7~UX\u00b7BSAT had the highest raw consistency (0.883) but was correctly "
    "excluded because its PRI (0.698) fell below 0.70.", align="justify")
add_table(
    [["Configuration", "UE", "UX", "BSAT", "Raw cov.", "Unique cov.", "Consistency"],
     ["C1", "\u25cf core", "\u25cf core", "\u25cf core", "0.497", "0.497", "0.857"],
     ["Solution", "", "", "", "0.497", "", "0.857"]],
    caption_above="Table 13. Sufficient configuration for high BSUC (Fiss notation).",
    note="\u25cf = core condition present.")
add_figure(f"{FIG}/Fig10_fsqca_XY_sufficiency.png",
           "Figure 10. fsQCA sufficiency XY plot for the core configuration UE\u00b7UX\u00b7BSAT "
           "(consistency = 0.857, coverage = 0.497).", 4.8)
add_figure(f"{FIG}/Fig12_truth_table.png",
           "Figure 12. Truth-table configurations ranked by sufficiency consistency "
           "(green = coded sufficient for high BSUC).", 5.5)

add_heading("4.4 Causal asymmetry (negated outcome)", 2)
add_para(
    "A separate sufficiency analysis for low continuance (~BSUC) returns the "
    "mirror-absence configuration ~UE\u00b7~UX\u00b7~BSAT (coverage = 0.492, "
    "consistency = 0.861). Because this is not a simple algebraic negation of the "
    "high-continuance recipe, the data demonstrate causal asymmetry: the conditions "
    "that produce continuance and those that produce abandonment are not perfect "
    "opposites. This asymmetric insight is unavailable from the symmetric PLS-SEM "
    "(Figure 13).", align="justify")
add_figure(f"{FIG}/Fig13_configurations.png",
           "Figure 13. Configurations for high vs. low BSUC (causal asymmetry; large circles = core conditions).", 4.8)

add_heading("4.5 Robustness of the fsQCA", 2)
add_para(
    "The configurational solution is stable under perturbations of the analytical "
    "choices (Table 14). Varying the calibration anchors (P90/P50/P10), raising the "
    "raw-consistency threshold to 0.85, and raising the frequency threshold to 6 "
    "all reproduce the identical UE\u00b7UX\u00b7BSAT solution with essentially unchanged "
    "coverage and consistency. Only fixed substantive anchors (6/4/2) simplify the "
    "recipe to UE\u00b7BSAT, consistent with UX being the weakest necessity condition.",
    align="justify")
add_table(
    [["Scenario", "# configs", "Sol. coverage", "Sol. consistency", "Configuration(s)"],
     ["Percentile 95/50/5 (main)", "1", "0.497", "0.857", "UE\u00b7UX\u00b7BSAT"],
     ["Percentile 90/50/10", "1", "0.476", "0.833", "UE\u00b7UX\u00b7BSAT"],
     ["Raw consistency \u2265 0.85", "1", "0.497", "0.857", "UE\u00b7UX\u00b7BSAT"],
     ["Frequency \u2265 6", "1", "0.497", "0.857", "UE\u00b7UX\u00b7BSAT"],
     ["Substantive anchors 6/4/2", "1", "0.533", "0.805", "UE\u00b7BSAT"]],
    caption_above="Table 14. fsQCA robustness checks.", size=10,
    col_align=["left","center","center","center","left"])

# =====================================================================
# 5. SYNTHESIS
# =====================================================================
add_heading("5. Synthesis: convergence and complementarity of PLS-SEM and fsQCA", 1)
add_para(
    "The two methods tell a coherent and mutually reinforcing story (Table 15). "
    "Convergence: both place satisfaction (BSAT) at the centre of continuance\u2014in "
    "PLS-SEM it is the strongest path and the full mediator of UE and UX; in fsQCA "
    "it has the highest single-condition necessity consistency and is a core "
    "element of the sole sufficient recipe. Both also agree that experience matters "
    "more than engagement upstream.", align="justify")
add_para(
    "Complementarity 1 (equifinality / conjunction): PLS-SEM finds the direct "
    "effects of UE and UX non-significant, but fsQCA clarifies that UE and UX are "
    "jointly required as part of the sufficient configuration\u2014the antecedents are "
    "complements, not substitutes. Complementarity 2 (asymmetry): fsQCA shows the "
    "recipe for low continuance is not the inverse of the recipe for high "
    "continuance, revealing causal asymmetry of managerial relevance.", align="justify")
add_table(
    [["Question", "PLS-SEM answer", "fsQCA answer"],
     ["What drives continuance on average?", "BSAT (\u03b2=.476***); UE/UX only indirectly", "BSAT highest necessity; core in recipe"],
     ["Are UE/UX effects direct?", "No \u2014 fully mediated by BSAT", "Required jointly, with BSAT"],
     ["Which combination yields high continuance?", "(not addressable)", "UE\u00b7UX\u00b7BSAT (cons.=.857, cov.=.497)"],
     ["Is low continuance the mirror image?", "(not addressable)", "No \u2014 ~UE\u00b7~UX\u00b7~BSAT (asymmetry)"],
     ["Predictive validity?", "Q\u00b2>0; PLS<LM on 3/4 indicators", "Solution stable across robustness checks"]],
    caption_above="Table 15. Summary of findings across methods.", size=10,
    col_align=["left","left","left"])
add_para(
    "In sum, satisfaction is the proximal engine of brand-system continuance, user "
    "experience is its dominant upstream antecedent, and the highest levels of "
    "continuance are reached only when engagement, experience and satisfaction are "
    "present in combination\u2014while the route to abandonment follows its own, "
    "asymmetric logic.", align="justify")

# ---------- Appendix figure ----------
add_heading("Appendix", 1)
add_figure(f"{FIG}/FigA1_item_correlations.png",
           "Figure A1. Item-level correlation matrix.", 5.5)

doc.save(OUT_DOCX)
print("Saved", OUT_DOCX)
