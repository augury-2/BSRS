"""
07_build_narrative_docx.py
Render the narrative (SDT / metaverse) Results section to Word. TNR 12 pt.
Prose-focused; references the numbered tables/figures in the main document.
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

OUT_DOCX = "Results_BSRS_narrative.docx"
doc = Document()


def set_font(style_name, size, bold=False, color=None):
    st = doc.styles[style_name]
    st.font.name = "Times New Roman"
    st.font.size = Pt(size)
    st.font.bold = bold
    rpr = st.element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts'); rpr.append(rfonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rfonts.set(qn(attr), "Times New Roman")
    if color is not None:
        st.font.color.rgb = color


set_font("Normal", 12)
for h, sz in [("Heading 1", 14), ("Heading 2", 13)]:
    set_font(h, sz, bold=True, color=RGBColor(0, 0, 0))
doc.styles["Normal"].paragraph_format.space_after = Pt(8)
doc.styles["Normal"].paragraph_format.line_spacing = 1.15


def para(text, italic=False, align="justify", shade=False, size=12):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.alignment = (WD_ALIGN_PARAGRAPH.JUSTIFY if align == "justify"
                   else WD_ALIGN_PARAGRAPH.LEFT)
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            r = p.add_run(part[2:-2]); r.bold = True
        else:
            r = p.add_run(part)
        r.font.name = "Times New Roman"; r.font.size = Pt(size); r.italic = italic
    if shade:
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), 'FFF2CC')
        pPr.append(shd)
    return p


def heading(text, level):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = "Times New Roman"; r.font.color.rgb = RGBColor(0, 0, 0)


def add_demo_table(rows, caption):
    from docx.enum.table import WD_TABLE_ALIGNMENT
    cap = doc.add_paragraph(); cap.paragraph_format.space_after = Pt(4)
    rr = cap.add_run(caption); rr.bold = True
    rr.font.name = "Times New Roman"; rr.font.size = Pt(11)
    n_cols = len(rows[0])
    t = doc.add_table(rows=0, cols=n_cols); t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    # category sub-header rows (those with empty n/% but non-empty label)
    cat_rows = set()
    for i, row in enumerate(rows):
        if i == 0:
            continue
        left_cat = row[0] and not row[1] and not row[2]
        right_cat = row[3] and not row[4] and not row[5]
        if left_cat or right_cat:
            cat_rows.add(i)
    for i, row in enumerate(rows):
        cells = t.add_row().cells
        for j in range(n_cols):
            cells[j].text = str(row[j])
            for p in cells[j].paragraphs:
                p.paragraph_format.space_after = Pt(1)
                p.paragraph_format.space_before = Pt(1)
                if j in (1, 2, 4, 5):
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.name = "Times New Roman"; r.font.size = Pt(10)
                    is_header = (i == 0)
                    # bold for header row and category labels
                    if is_header:
                        r.bold = True
                    elif i in cat_rows and j in (0, 3):
                        r.bold = True
            if i == 0:
                tcPr = cells[j]._tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), 'DDDDDD')
                tcPr.append(shd)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return t


# Title
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Results"); r.bold = True
r.font.name = "Times New Roman"; r.font.size = Pt(16)
t.paragraph_format.space_after = Pt(12)

# ---- Sample characteristics ----
heading("Sample characteristics", 1)
para("The final sample comprised 312 complete responses from active users of "
     "avatar-mediated (metaverse-type) platforms in India; all 312 cases were "
     "retained, as there were no missing values across the 20 items and no "
     "straight-line response patterns. The demographic and platform-usage profile of "
     "the sample is summarised in Table 1.")
para("The sample was reasonably balanced by gender, with a slight majority of women "
     "(54.00%, n = 169) relative to men (46.00%, n = 143). Respondents were "
     "distributed fairly evenly across the five age bands, from 18\u201322 years (20.51%) "
     "and 23\u201328 years (24.04%) through to 42\u201345 years (17.95%), indicating that the "
     "sample spans younger and more established adult cohorts rather than being "
     "concentrated in a single generation. In terms of marital status, the largest "
     "groups were married respondents without children (38.78%) and single "
     "respondents (34.29%), followed by married respondents with children (26.93%). "
     "Occupationally, the sample combined business professionals (39.10%), salaried "
     "employees (32.69%), and students (28.21%), and monthly family income was spread "
     "across all bands, with the largest group in the INR 50,001\u201380,000 range "
     "(32.05%) and meaningful representation at both lower (\u2264 INR 30,000: 18.40%) and "
     "higher (\u2265 INR 80,001: 23.91%) ends.")
para("Crucially for the present study, respondents reported substantial experience "
     "with brand-relevant activities in immersive environments. A clear majority "
     "engaged with metaverse-type platforms at least weekly (71.46%: daily 16.67%, "
     "several times a week 31.41%, weekly 23.38%), and the remainder did so monthly "
     "(17.31%) or rarely (11.21%). Beyond mere access, large proportions of "
     "respondents had participated in behaviours central to brand experience in these "
     "settings: social interaction and content creation (52.56%), NFT interaction "
     "(45.19%), and virtual-event participation (41.35%). Together, this profile "
     "confirms that the respondents possessed meaningful, behaviourally grounded "
     "exposure to brands in avatar-mediated environments, supporting the relevance of "
     "the sample to the research questions.")

demo_rows = [
    ["Participants' information", "n", "%", "Participants' information", "n", "%"],
    ["Gender", "", "", "Metaverse engagement frequency", "", ""],
    ["Male", "143", "46.00", "Daily", "52", "16.67"],
    ["Female", "169", "54.00", "Several times a week", "98", "31.41"],
    ["", "", "", "Weekly", "73", "23.38"],
    ["Age (years)", "", "", "Monthly", "54", "17.31"],
    ["18\u201322", "64", "20.51", "Rarely", "35", "11.21"],
    ["23\u201328", "75", "24.04", "", "", ""],
    ["29\u201334", "60", "19.23", "NFT interaction", "", ""],
    ["35\u201341", "57", "18.27", "Yes", "141", "45.19"],
    ["42\u201345", "56", "17.95", "No", "171", "54.81"],
    ["Marital status", "", "", "Virtual event participation", "", ""],
    ["Single", "107", "34.29", "Yes", "129", "41.35"],
    ["Married with children", "84", "26.93", "No", "183", "58.65"],
    ["Married without children", "121", "38.78", "Social interaction / content creation", "", ""],
    ["", "", "", "Yes", "164", "52.56"],
    ["Occupation", "", "", "No", "148", "47.44"],
    ["Students", "88", "28.21", "Monthly family income", "", ""],
    ["Job", "102", "32.69", "INR 30,000 or less", "57", "18.40"],
    ["Business", "122", "39.10", "INR 30,001\u201350,000", "80", "25.64"],
    ["", "", "", "INR 50,001\u201380,000", "100", "32.05"],
    ["", "", "", "INR 80,001 and above", "75", "23.91"],
]
tbl = add_demo_table(demo_rows, "Table 1. Sample profile (N = 312).")

# ---- Measurement model ----
heading("Measurement model", 1)
para("We first assessed the reliability and validity of the reflective measurement "
     "model using PLS-SEM (Mode A; path-weighting scheme). Standardized factor "
     "loadings for all indicators were positive and statistically significant, "
     "ranging from 0.812 to 0.879\u2014comfortably above the 0.708 benchmark and well "
     "clear of the more lenient 0.66 floor\u2014indicating strong indicator reliability "
     "(Table 2; Figure 7). Internal consistency was satisfactory and non-redundant: "
     "Cronbach's alpha ranged from 0.885 to 0.892, the Dijkstra\u2013Henseler \u03c1_A from "
     "0.887 to 0.895, and composite reliability (\u03c1_c) from 0.918 to 0.922, all "
     "comfortably exceeding the 0.70 threshold while remaining below 0.95. Average "
     "variance extracted (AVE) values ranged from 0.692 to 0.747, exceeding the 0.50 "
     "guideline and indicating that each construct captures a substantial majority of "
     "the variance in its indicators relative to measurement error.")
para("Discriminant validity was evaluated using the heterotrait\u2013monotrait (HTMT) "
     "ratio of correlations. All HTMT values among UE, UX, BSAT, and BSUC were below "
     "the conservative 0.85 cut-off (range 0.30\u20130.62), with the highest value "
     "observed for the BSAT\u2013BSUC pair (0.615) and the lowest for the UE\u2013BSUC pair "
     "(0.298), indicating that each construct is empirically distinguishable from the "
     "others (Table 4; Figure 5). The Fornell\u2013Larcker criterion corroborated this: "
     "for every construct, the square root of the AVE exceeded its correlations with "
     "all other constructs, and the inter-construct correlations were moderate "
     "(0.27\u20130.55) and below the respective reliability estimates (Table 3; "
     "Figure 6). The cross-loading matrix showed each indicator loading highest on "
     "its intended construct by a margin of at least 0.40.")
para("To assess common method bias, we complemented procedural remedies with "
     "statistical diagnostics. Harman's single-factor test attributed only 34.4% of "
     "the variance to the first unrotated factor, below the 50% threshold. A full "
     "collinearity assessment confirmed this: variance inflation factor (VIF) values "
     "for all latent variables were well below the recommended 3.3 cut-off (maximum "
     "inner VIF = 1.35; maximum outer VIF = 2.46), suggesting that common method "
     "variance is unlikely to pose a serious threat (Kock, 2015). Taken together, "
     "these results support the adequacy of the measurement model.")

# ---- Structural model ----
heading("Structural model", 1)
para("We next examined the structural model to test H1\u2013H3, assessing path "
     "significance via bootstrapping with 5,000 resamples. Because H1\u2013H3 are "
     "directional hypotheses, one-tailed tests were used for the hypothesised paths, "
     "consistent with standard PLS-SEM practice (Hair et al., 2022). The model "
     "explained meaningful variance in the endogenous constructs (R\u00b2 = 0.224 for "
     "BSAT and R\u00b2 = 0.315 for BSUC) and demonstrated predictive relevance "
     "(Q\u00b2 = 0.160 and 0.228, respectively; SRMR = 0.047), with out-of-sample "
     "validity confirmed by PLSpredict.")
para("User engagement (UE) and user experience (UX) both showed positive, "
     "statistically significant direct relationships with brand success (BSUC) "
     "(UE \u2192 BSUC: \u03b2 = 0.088, t = 1.82, one-tailed p = .035; UX \u2192 BSUC: \u03b2 = 0.095, "
     "t = 1.78, one-tailed p = .038), supporting H1 and H2 (Table 8; Figure 1). Users "
     "who reported higher cognitive, emotional, and behavioural involvement with "
     "brands in avatar-mediated environments, and who rated their interaction "
     "experiences as more immersive, enjoyable, and usable, also reported stronger "
     "continuance, purchase, and advocacy intentions toward those brands.")
para("Brand satisfaction (BSAT) was strongly and positively associated with BSUC "
     "(\u03b2 = 0.476, t = 10.15, p < .001), and both UE and UX were positively associated "
     "with BSAT (UE \u2192 BSAT: \u03b2 = 0.157, t = 2.89, p = .004; UX \u2192 BSAT: \u03b2 = 0.393, "
     "t = 7.59, p < .001), indicating that more engaging and higher-quality "
     "avatar-mediated experiences translate into more favourable overall evaluations "
     "of the brand. Bootstrapped indirect effects from UE and UX to BSUC via BSAT "
     "were both positive and significant (UE \u2192 BSAT \u2192 BSUC: 0.075, 95% CI "
     "[0.024, 0.132]; UX \u2192 BSAT \u2192 BSUC: 0.187, 95% CI [0.131, 0.251]), supporting "
     "H3 and indicating that brand satisfaction mediates the influence of engagement "
     "and experience on brand success (Table 9). Because the direct effects of UE and "
     "UX on BSUC remained significant when BSAT was included, and because each direct "
     "effect and its corresponding indirect effect carry the same (positive) sign, "
     "the data point to complementary partial mediation (Zhao et al., 2010; Nitzl et "
     "al., 2016) rather than full mediation. The total effects on brand success "
     "ranked the antecedents as BSAT (0.476) > UX (0.282) > UE (0.162).")
para("Overall, the structural model provides evidence that SDT-consistent "
     "manifestations of need satisfaction at the interaction level\u2014captured through "
     "UE and UX\u2014shape brand-level outcomes both directly and through users' "
     "satisfaction with the brand's immersive presence.")

# ---- fsQCA ----
heading("Configurational analysis (fsQCA)", 1)
para("To explore whether different combinations of UE, UX, and BSAT can lead to high "
     "brand success, we applied fuzzy-set qualitative comparative analysis. Composite "
     "scores for the three conditions (UE, UX, BSAT) and the outcome (BSUC) were "
     "calibrated into fuzzy sets using the direct method with theoretically informed, "
     "percentile-based anchors\u2014full membership at the 95th percentile, the cross-over "
     "at the 50th percentile, and full non-membership at the 5th percentile (Table 10; "
     "Figure 9). The resulting truth table was analysed using standard thresholds "
     "(frequency \u2265 4 cases; raw consistency \u2265 0.80; PRI \u2265 0.70) to identify "
     "configurations sufficient for high BSUC (Table 12).")
para("Analysis of necessary conditions. The fsQCA results show that no single "
     "condition is necessary for high brand success: high UE, high UX, and high BSAT "
     "each fell below the 0.90 necessity-consistency threshold (maximum = 0.751 for "
     "BSAT), consistent with the SDT-based view that different patterns of need "
     "satisfaction can sustain favourable outcomes (Table 11; Figure 11).")
para("Analysis of sufficient configurations. The minimisation identified a single, "
     "highly consistent sufficient configuration for high brand success\u2014the "
     "simultaneous presence of high UE, high UX, and high BSAT (raw consistency = "
     "0.857, PRI = 0.737, raw coverage = 0.497; Table 13b; Figure 10). All three "
     "conditions are core because they appear in both the parsimonious and the "
     "intermediate solution (Fiss, 2011). This \u201cfull recipe\u201d captures scenarios in "
     "which avatar-mediated environments strongly satisfy autonomy, competence, and "
     "relatedness needs, producing high engagement, rich experiences, and strong "
     "satisfaction and, consequently, the strongest brand-relevant outcomes.")
para("An extended subset\u2013superset analysis (Table 13a) clarifies this result by "
     "ordering all candidate sufficient terms. It reveals the canonical "
     "consistency\u2013coverage trade-off: the triadic term UE \u00b7 UX \u00b7 BSAT attains the "
     "highest solution consistency (0.857), while less restrictive supersets achieve "
     "broader coverage at the cost of consistency\u2014satisfaction alone (BSAT) reaches "
     "the highest coverage (0.751) but a sub-threshold PRI (0.661), and the "
     "two-condition terms UE \u00b7 BSAT (consistency 0.839) and UX \u00b7 BSAT (consistency "
     "0.829) sit in between. Thus, although satisfaction is the empirically dominant "
     "ingredient, it is the conjunction of engagement, experience, and satisfaction "
     "that constitutes the most trustworthy recipe for high brand success.")
para("Causal asymmetry. A separate analysis of the negated outcome returned the "
     "mirror-absence configuration ~UE \u00b7 ~UX \u00b7 ~BSAT (consistency = 0.861, "
     "coverage = 0.492; Figure 13), demonstrating that the recipe for low brand "
     "success is not the simple inverse of the recipe for high success.")
para("Taken together, the configurational analysis complements the PLS-SEM results "
     "by showing that, while brand success in avatar-mediated environments is driven "
     "proximally by satisfaction, the highest and most consistent levels of success "
     "arise from the joint, SDT-consistent combination of engagement, experience, and "
     "satisfaction operating together, with abandonment following its own asymmetric "
     "logic.")

para("Divergence from the supplied draft narrative (fsQCA). The draft provided "
     "described three sufficient configurations (BSAT\u00b7UX with moderate UE; BSAT\u00b7UE; "
     "and UE\u00b7UX in the absence of high BSAT). The truth table derived from "
     "Responses.xlsx does not reproduce these under standard thresholds: the "
     "UE\u00b7UX\u00b7~BSAT corner has raw consistency 0.764 (< 0.80) and PRI 0.443 (< 0.70), "
     "so it is not sufficient; and minimising the BSAT-present corners collapses them "
     "to a single term. Under any defensible threshold set, the data yield one core "
     "\u201cfull-recipe\u201d configuration (UE\u00b7UX\u00b7BSAT). If a multi-configuration solution is "
     "required, it would need to come from a different dataset/run or from "
     "deliberately relaxed thresholds (which would lower solution validity and should "
     "be justified explicitly). The result reported here is the one the current data "
     "actually support.", italic=True, shade=True)

doc.save(OUT_DOCX)
print("Saved", OUT_DOCX)
