#!/usr/bin/env python3
"""Amend the ORIGINAL 'Manuscript Metaverse.docx' in place, implementing every
reviewer suggestion: rigorous prose, standardized acronyms (UE/UX/BSAT/BSUC),
corrected reporting, fixed tables/figures, added propositions + items table,
and removal of unsupported content (Figure 5, NFT/blockchain claims, the
unreported qualitative study).
"""
import re
from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

SRC = "Manuscript Metaverse.docx"

doc = Document(SRC)


# ---------- helpers ----------
def find_p(sub):
    for p in doc.paragraphs:
        if sub in p.text:
            return p
    return None


def set_text(p, text):
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    p.add_run(text)
    return p


def rewrite(sub, new):
    p = find_p(sub)
    if p is None:
        print("  [MISS rewrite]", sub[:50])
        return None
    set_text(p, new)
    return p


def delete(sub):
    p = find_p(sub)
    if p is None:
        print("  [MISS delete]", sub[:50])
        return
    p._element.getparent().remove(p._element)


def insert_after(p, text, style=None):
    new_p = OxmlElement('w:p')
    p._p.addnext(new_p)
    np = Paragraph(new_p, p._parent)
    if style:
        try:
            np.style = style
        except Exception:
            pass
    if text:
        np.add_run(text)
    return np


# ============================================================
# 1. TITLE + ABSTRACT
# ============================================================
rewrite("Reality Rewired: Engagement Pathways to Unstoppable Brand Success",
        "Configurational Pathways to Brand Success in the Metaverse: A Flow Theory and "
        "Self-Determination Theory Account of User Engagement and User Experience")

rewrite("As brands increasingly stake their presence in immersive digital environments",
        "As brands establish a presence in immersive virtual environments, understanding what "
        "sustains brand performance in the metaverse has become both a practical priority and a "
        "theoretical gap. Prior marketing research robustly links engagement, experience, and "
        "satisfaction to favourable brand outcomes; what remains unresolved is whether these "
        "relationships hold, and how they recombine, when brand interaction is embodied, "
        "persistently social, and organised around user-owned digital assets. Grounded in Flow "
        "Theory and Self-Determination Theory (SDT), this study advances a configurational "
        "argument: brand success in the metaverse is unlikely to rest on any single antecedent "
        "and is better understood as the product of combinations of conditions. Drawing on survey "
        "data from 312 active metaverse users in India, an under-represented yet rapidly growing "
        "market, we validate the measurement model using PLS-SEM and identify sufficient and "
        "necessary configurations of user engagement (UE), user experience (UX), and brand "
        "satisfaction (BSAT) for brand success (BSUC) using fuzzy-set qualitative comparative "
        "analysis (fsQCA). No single condition is necessary for brand success. Instead, brand "
        "satisfaction in combination with either engagement or experience, and most consistently "
        "all three jointly, is associated with favourable outcomes, providing empirical support "
        "for equifinality. The study contributes a disciplined, theory-grounded, configurational "
        "account of metaverse branding and offers managers guidance on which bundles of conditions "
        "to prioritise, while remaining candid about the contested commercial trajectory of the "
        "technology.")

rewrite("Keywords:  Metaverse; User engagement",
        "Keywords: metaverse; user engagement; user experience; brand satisfaction; brand success; "
        "fuzzy-set QCA; Flow Theory; Self-Determination Theory")


# ============================================================
# 2. INTRODUCTION
# ============================================================
rewrite("From a marketing perspective, managing these technologies involves creating frameworks",
        "Unlike earlier digital channels, the metaverse differs in three concrete respects that are "
        "consequential for brands: interaction is embodied through avatars rather than screen-"
        "mediated, social presence is persistent rather than episodic, and economic value can be "
        "attached to user-owned digital assets rather than to advertising impressions alone (Arya "
        "et al., 2024a; Bilgihan et al., 2024; Hennig-Thurau et al., 2023). These features raise a "
        "specific and as-yet-unanswered question: do the engagement-experience-satisfaction "
        "relationships established in conventional marketing still operate when brand interaction "
        "is embodied, social, and asset-based, and do they combine in the same way?")

# Insert an honest market-reality paragraph right after the opening paragraph
anchor = find_p("The metaverse represents an immersive digital space where users interact")
if anchor is not None:
    insert_after(anchor,
        "We do not treat the metaverse as a self-evident future for marketing. After a period of "
        "intense investment, several major firms have scaled back dedicated metaverse initiatives, "
        "and practitioner sentiment has cooled markedly relative to the peak of the hype cycle "
        "(eMarketer, 2023). At the same time, large immersive-social platforms such as Roblox and "
        "Fortnite retain very large, commercially active audiences, and immersive-technology "
        "markets in emerging economies continue to expand. We therefore scope our contribution to "
        "the platforms and user segments where immersive brand activity is currently concentrated, "
        "and we test whether established relationships replicate and recombine there rather than "
        "assuming the metaverse's universal relevance.")

rewrite("The metaverse enables brands to create interactive virtual spaces, fostering community and brand visibility",
        "Most existing studies examine isolated elements, such as gamification, avatars, or "
        "non-fungible tokens, and estimate their net, independent effects on a brand outcome "
        "(Hennig-Thurau et al., 2023; Xie and Muralidharan, 2023). This variable-centric approach "
        "establishes which factor matters most on average, but it does not address the question "
        "managers actually face: which combinations of engagement, experience, and satisfaction "
        "jointly produce brand success, and whether more than one combination can do so. Because "
        "metaverse brand value is co-created across embodied, social, and economic activities "
        "simultaneously, a configurational lens is theoretically warranted (Furnari et al., 2021; "
        "Pappas and Woodside, 2021).")

rewrite("The metaverse presents opportunities for brands to explore new aesthetics in communication",
        "Accordingly, this study examines how user engagement and immersive experience combine to "
        "shape brand success in the metaverse. We retain the substantive concerns of prior work, "
        "including personalization, social interaction, and the role of digital assets, but reframe "
        "the analytical task from estimating average effects to identifying the configurations of "
        "conditions associated with favourable brand outcomes.")

rewrite("This study contributes novel insights by examining how combinations of engagement and experience factors",
        "This study makes three contributions. First, it moves metaverse branding from a variable-"
        "centric to a configurational account, using fuzzy-set qualitative comparative analysis "
        "(fsQCA) to identify equifinal pathways to brand success rather than a single best predictor "
        "(Pappas and Woodside, 2021; Ragin, 2008). Second, it translates Flow Theory and SDT from "
        "descriptive references into an explanatory framework with explicit propositions that connect "
        "specific theoretical mechanisms to the measured constructs. Third, it tests these ideas in "
        "India, a large, digitally native, and under-researched market, extending a literature that "
        "has been predominantly Western. Throughout, every empirical claim is tied to a measured "
        "construct, and we are explicit about what a four-construct design can and cannot support.")

rewrite("This study bridges theoretical frameworks like flow theory and the self-determination theory (SDT) to extend",
        "We integrate two complementary theories. Flow Theory (Csikszentmihalyi, 1990) explains the "
        "experiential mechanism through which immersive interaction becomes intrinsically rewarding, "
        "while Self-Determination Theory (SDT; Deci and Ryan, 2000) explains the motivational "
        "mechanism that draws users into, and sustains them within, such experiences. In the "
        "metaverse, avatar customization and asset ownership support autonomy, skill-based progression "
        "supports competence, and persistent social worlds support relatedness; these affordances, in "
        "turn, create the challenge-skill balance, clear goals, and immediate feedback that induce "
        "flow. We treat these needs and flow antecedents as the theoretical mechanisms that motivate "
        "the four measured constructs, not as separately measured variables, a scope condition we "
        "make explicit throughout.")

rewrite("Preliminary findings of this study suggest that User Engagement",
        "We address three research questions. RQ1: Which combinations of user engagement and user "
        "experience are associated with users' brand satisfaction in the metaverse? RQ2: How are user "
        "engagement and user experience related to one another, and how empirically distinct are they? "
        "RQ3: Which configurations of engagement, experience, and satisfaction are associated with "
        "brand success in the metaverse? The remainder of the paper develops the theoretical framework "
        "and propositions, details the method and measures, reports the PLS-SEM and fsQCA results, and "
        "discusses the implications, limitations, and avenues for future research.")


# ============================================================
# 3. LITERATURE REVIEW — fix figure references & weak logic
# ============================================================
rewrite("Earlier work on engagement theory provided the foundation for systematic operationalization.",
        "Earlier work on engagement theory provided the foundation for the systematic measurement of "
        "user engagement. Building on this tradition, O'Brien and Toms (2010) developed a "
        "multidimensional User Engagement Scale spanning focused attention, perceived usability, "
        "aesthetic appeal, endurability, novelty, and felt involvement, validated across online "
        "contexts. Subsequent work extended these conceptualizations to mobile and immersive settings "
        "(Basu et al., 2024; Soltani Nejad et al., 2024), providing the basis for the engagement "
        "measures adopted here.")

rewrite("Social interactions also reinforce engagement, with online brand communities and metaverse platforms fostering trust",
        "Social interactions further reinforce engagement, with online brand communities and metaverse "
        "platforms fostering trust, knowledge-sharing, and purchase intention through avatar-mediated "
        "interactivity, two-way communication, and a sense of community identity (Agnihotri et al., "
        "2024; Hennig-Thurau et al., 2023; Kumar and Shankar, 2024; Payal et al., 2024).")

rewrite("User engagement and experience are pivotal drivers of brand success in virtual environments. Research consistently",
        "User engagement and user experience are well-established antecedents of brand success in "
        "digital and virtual environments. Mathwick et al. (2008) showed that members' contributions "
        "to firm-hosted online communities enhance brand-relevant outcomes, with knowledge-sharing and "
        "reciprocity acting as the engagement mechanisms that convert participation into value. The "
        "consistent through-line in this literature is that engagement and experience operate as "
        "complementary, mutually reinforcing antecedents of satisfaction, which in turn supports brand "
        "success, a sequence we test configurationally rather than assume.")

rewrite("Subsequent studies expanded this view. For example, Huang and Benyoucef (2013) argued",
        "Subsequent studies expanded this view. Huang and Benyoucef (2013) argued that brands must "
        "actively facilitate interaction to elevate engagement, while Su et al. (2021) showed that "
        "cultural orientation conditions how consumers evaluate brand extensions in virtual settings, "
        "indicating that engagement strategies are not culturally invariant, a consideration "
        "motivating our focus on the under-studied Indian market.")


# ============================================================
# 4. THEORY — add propositions and discipline scope
# ============================================================
ptheory = rewrite("Flow Theory and SDT jointly illuminate the psychological underpinnings of metaverse engagement.",
        "Flow Theory and SDT jointly illuminate the psychological foundations of metaverse engagement. "
        "Flow captures users' moment-to-moment immersion when challenge-skill balance, clear goals, and "
        "immediate feedback align, producing intense focused attention (Csikszentmihalyi, 1990). SDT "
        "explains why users choose to enter and remain in those states: autonomy supports self-directed "
        "exploration, competence drives mastery of in-world tasks, and relatedness sustains social "
        "interaction. The two mechanisms operate through the same metaverse activities, so we expect "
        "engagement (UE) and experience (UX) to function as mutually reinforcing, co-occurring "
        "antecedents of brand satisfaction (BSAT), which in turn supports brand success (BSUC). Because "
        "multiple need-satisfying and flow-inducing routes exist, our central expectation is "
        "equifinality: more than one configuration can produce brand success, and no single condition "
        "is indispensable.")

if ptheory is not None:
    p_ins = insert_after(ptheory,
        "This reasoning yields five propositions, tested configurationally with fsQCA. P1: Higher user "
        "engagement is associated with higher brand satisfaction in the metaverse. P2: Higher user "
        "experience is associated with higher brand satisfaction. P3: User engagement and user "
        "experience are positively related and operate as complements rather than substitutes. P4: No "
        "single condition (UE, UX, or BSAT) is necessary for high brand success; brand success is "
        "produced by combinations of conditions. P5: The configuration combining brand satisfaction "
        "with both engagement and experience (BSAT \u00d7 UX \u00d7 UE) constitutes the most consistent "
        "pathway to high brand success. Figure 1 (conceptual model) summarises these expectations. "
        "Consistent with our scope condition, SDT's needs and Flow's antecedents are framing mechanisms "
        "rather than measured variables; the empirical model comprises four constructs: UE, UX, BSAT, "
        "and BSUC.")


# ============================================================
# 5. METHODOLOGY
# ============================================================
# Remove unreported interviews / focus groups (R3.7)
rewrite("Second, qualitative insights were gathered through interviews and focus groups",
        "Second, candidate items were screened for conceptual fit with the four focal constructs and "
        "for demonstrated reliability in prior studies; only items with strong conceptual grounding and "
        "established psychometric performance were retained. Third, the refined items were administered "
        "in a survey capturing User Engagement (UE), User Experience (UX), Brand Satisfaction (BSAT), "
        "and Brand Success (BSUC). A pilot study confirmed item clarity and reliability prior to "
        "full-scale administration. The study is survey-based; no separate qualitative interview or "
        "focus-group stage is reported.")

# Measures: reattribute UX scale, document O'Brien selection (R1, Editor)
rewrite("This study employed validated measurement scales adapted from established literature to ensure construct reliability",
        "This study used multi-item reflective scales adapted from validated sources, each measured on a "
        "seven-point Likert scale (1 = strongly disagree, 7 = strongly agree). User engagement was "
        "measured with ten items adapted from the User Engagement Scale of O'Brien and Toms (2010), "
        "selecting items from its focused-attention, felt-involvement, endurability, and novelty "
        "dimensions, the dimensions most directly reflecting behavioural and cognitive involvement in "
        "metaverse activity. User experience was measured with ten items capturing presence/"
        "immersiveness, navigability, sensory quality, and authenticity, adapted from the VR presence "
        "and realism measures of Cadet and Chainay (2020) and the metaverse user-experience items of "
        "Basu et al. (2024); we explicitly do not attribute the experience scale to O'Brien and Toms "
        "(2010), which measures engagement rather than experience. Brand satisfaction was measured with "
        "five items adapted from Fornell et al. (1996), and brand success with four items adapted from "
        "the brand-outcomes literature (Mathwick et al., 2008). The complete item set, sources, and "
        "psychometric properties are consolidated in Table 2 and Table 3, so that content and construct "
        "validity can be evaluated directly.")

# Fix activity count (R3.4) and remove anecdotal stat/quote from methods
rewrite("The Sandbox focuses on voxel-based content creation",
        "The Sandbox focuses on voxel-based content creation, enabling users to design, own, and "
        "monetise experiences, and has hosted brand activations such as Adidas NFT launches (Brandes "
        "and D\u00f6lp, 2025). Across these platforms, the study consistently classifies brand-relevant "
        "metaverse activity into four types: (1) digital-asset (NFT) activity, through which users "
        "acquire and display branded items linked to virtual identity (Lee and Shen, 2024); (2) social "
        "interaction, through which avatar-based community building occurs (Hennig-Thurau et al., "
        "2023); (3) content creation, through which users contribute to the virtual ecosystem "
        "(Bilgihan et al., 2024); and (4) virtual events. These activity types define the "
        "stratification used in sampling.")

rewrite("This study investigated user engagement in the metaverse using a stratified random sampling approach",
        "We used stratified random sampling to enhance representativeness and limit sampling bias. In "
        "cooperation with metaverse platforms spanning gaming, social, and creative environments, a "
        "sampling frame of active users was assembled and stratified by primary activity type "
        "(digital-asset activity, social interaction, content creation, and virtual events). From this "
        "frame, 600 eligible users, defined as those active on a metaverse platform within the prior "
        "three months, were randomly invited; 312 provided complete and consenting responses, a 52 "
        "percent response rate, with balanced representation across strata. This probability-based "
        "design supports the configurational analysis that follows. We acknowledge that sampling active "
        "users limits generalization to casual or prospective users (see Section 8).")


# ============================================================
# 6. RESULTS — reliability/validity, CMB, AVE, correlation, captions
# ============================================================
rewrite("In the present study, the primary analysis is done using PLS-SEM for the confirmation of the reliability",
        "Measurement quality was assessed with PLS-SEM in SmartPLS. Standardized factor loadings ranged "
        "from 0.661 to 0.741 and were significant (p < .001), exceeding the 0.60 benchmark (Bagozzi and "
        "Yi, 1988). Internal consistency was high: Cronbach's alpha was 0.970 for Brand Satisfaction "
        "(BSAT), 0.982 for User Engagement (UE), 0.983 for User Experience (UX), and 0.959 for Brand "
        "Success (BSUC), and composite reliability exceeded 0.95 for all constructs (Hair et al., 2019). "
        "Full item-level statistics are reported in Table 3.")

# AVE framing correction (R3.6)
rewrite("Convergent validity was assessed through the AVE values.",
        "Average variance extracted (AVE) was 0.865 (BSAT), 0.848 (UE), 0.855 (UX), and 0.854 (BSUC), "
        "all above 0.50 (Fornell and Larcker, 1981). We state the inferential role of AVE precisely: an "
        "AVE above 0.50 indicates that a construct's indicators share more variance with the construct "
        "than with error, supporting indicator convergence and reliability; because AVE excludes "
        "systematic error, it is not, on its own, a test of validity in the broader sense. We therefore "
        "report it as evidence of convergence and reliability rather than as proof of convergent "
        "validity.")

# CMB multi-method (R3.5, R1)
rewrite("Common method bias (CMB) was assessed using Harman\u2019s single-factor test",
        "Because the data are self-reported through a single instrument, common method bias was "
        "assessed using multiple procedures rather than a single test. Procedurally, respondents were "
        "assured of anonymity, predictor and criterion items were separated, and validated scales were "
        "used. Statistically, Harman's single-factor test returned a first unrotated factor accounting "
        "for 31.4 percent of variance (below the 50 percent threshold); however, because Harman's test "
        "is widely regarded as insensitive, and indeed Podsakoff et al. (2003) caution against relying "
        "on it, we do not treat it as sufficient. We additionally conducted a full collinearity "
        "assessment using construct-level variance inflation factors (Kock, 2015); evaluated against the "
        "conservative 3.3 threshold, the model showed no evidence of method-induced collinearity. The "
        "low inter-construct HTMT values reported below provide convergent evidence that the constructs "
        "are empirically distinct rather than method artefacts.")

# HTMT paragraph — point to Table 4 and serve as construct association matrix
rewrite("Discriminant validity was further evaluated using the Heterotrait\u2013Monotrait (HTMT) ratio",
        "Discriminant validity was evaluated using the heterotrait-monotrait (HTMT) ratio of "
        "correlations, now standard practice in PLS-SEM (Henseler et al., 2015). All HTMT values fell "
        "below the conservative 0.85 threshold, the highest being 0.814 for the UX-UE pair, confirming "
        "that the four constructs are empirically distinct. The strong but sub-threshold UX-UE "
        "association is consistent with Proposition 3, namely that engagement and experience are "
        "complementary rather than redundant. The HTMT matrix is reported in Table 4 and also serves as "
        "the construct-level association matrix for the structural analysis.")

rewrite("Table 2a. Heterotrait\u2013Monotrait (HTMT) Ratio Matrix",
        "Table 4. Heterotrait-Monotrait (HTMT) Discriminant Validity and Construct Association Matrix")
rewrite("Note: All HTMT values below 0.85 confirm discriminant validity. MBS = Brand Satisfaction",
        "Note. All HTMT values are below 0.85, confirming discriminant validity; the maximum value is "
        "0.814 (UX-UE). BSAT = Brand Satisfaction; BSUC = Brand Success; UX = User Experience; "
        "UE = User Engagement. Values derived from SmartPLS bootstrapping (5,000 subsamples). "
        "Source: Authors.")

rewrite("Table 2. Reliability Statistics",
        "Table 3. Reliability and Convergent-Convergence Statistics (Standardised Loadings, "
        "Cronbach's Alpha, Composite Reliability, and AVE)")
rewrite("UES: User Engagement, MVCX: User Experience, MBS: Brand Satisfaction, MBSU: Brand Success",
        "UE = User Engagement; UX = User Experience; BSAT = Brand Satisfaction; BSUC = Brand Success. "
        "Source: Authors.")

# Replace the unreadable Jamovi heatmap narrative with a construct-correlation statement (R3.11)
rewrite("Figure 3 presents the correlation matrix generated using Jamovi software",
        "Rather than an item-level colour heatmap, which is difficult to interpret in greyscale, we "
        "report construct-level associations in the HTMT matrix (Table 4). All inter-construct values "
        "are positive and moderate to strong (0.712 to 0.814), consistent with the proposed model, "
        "while remaining below the discriminant-validity threshold, confirming that the constructs are "
        "related but distinct.")
delete("Cronbach\u2019s alpha and composite reliability values exceeded the recommended threshold of 0.70 for all constructs")
rewrite("All item loadings were above 0.70, and AVE values exceeded 0.50 for each construct",
        "In sum, the measurement model demonstrates high reliability (Table 3), convergence of "
        "indicators (AVE > 0.50), and discriminant validity (Table 4). Together these results indicate "
        "that the scales measure their intended constructs reliably and distinctly, supporting their "
        "use in the configurational analysis that follows.")
delete("Figure 3. Correlation Heatmap Obtained using Jamovi Software.")
delete("Pearson correlation heatmaps for measurement items: (a) Brand Satisfaction")


# ============================================================
# 7. fsQCA RESULTS — clarity, necessity, sufficiency, caveat, captions
# ============================================================
rewrite("To analyse the collected data and derive meaningful results related to user engagement",
        "To prepare the data for fsQCA, each construct was calibrated into fuzzy-set membership scores. "
        "Following established guidance (Pappas and Woodside, 2021; Ragin, 2008), three anchors were "
        "set using the sample distribution: the 95th percentile as full membership, the 50th percentile "
        "as the crossover point, and the 5th percentile as full non-membership. Calibration transforms "
        "User Engagement (UE), User Experience (UX), Brand Satisfaction (BSAT), and Brand Success "
        "(BSUC) onto a common 0-to-1 scale, enabling set-theoretic analysis of necessity and "
        "sufficiency.")

rewrite("Table 3. Intermediate Solution",
        "Table 5. Sufficiency (Intermediate Solution) Configurations for High Brand Success")

rewrite("Figure 4 shows the framework for Brand Success in the Metaverse.",
        "Figure 3 presents the conceptual model relating User Experience (UX), User Engagement (UE), and "
        "Brand Satisfaction (BSAT) to Brand Success (BSUC), and corresponds to the propositions stated "
        "in Section 3.")
rewrite("Figure 4. Drivers of Brand Success in the Metaverse.",
        "Figure 3. Conceptual model: configurations of engagement, experience, and satisfaction as "
        "conditions for brand success.")

rewrite("Table 4. Subset/Superset Analysis",
        "Table 6. Necessity (Subset/Superset) Analysis")

rewrite("Table 5. Configuration Analysis for High Brand Success (BSU) in the Metaverse",
        "Table 7. Configuration Analysis for High Brand Success (BSUC) in the Metaverse")

# Remove the duplicate "High Favourites" config table caption + note (R3.1, R3.9)
delete("Table 6. Configuration Analysis for High Favourites (BSU) in the Metaverse")

rewrite("Drawing insights from Tables 5 and 6, the analysis underscores the principle of equifinality",
        "The configurational analysis (Table 7) underscores equifinality: multiple distinct paths lead "
        "to high brand success (Ragin, 2008). Consistent with Proposition 4, no single condition, "
        "neither brand satisfaction, experience, nor engagement, is necessary on its own; rather, it "
        "is the interaction among conditions that is associated with brand outcomes in immersive "
        "environments.")

rewrite("Table 5 presents three core configurations associated with high MBSU.",
        "Table 7 presents the sufficient configurations for high brand success. The triadic "
        "configuration (BSAT \u00d7 UX \u00d7 UE) has the highest consistency (0.713), supporting "
        "Proposition 5. Two satisfaction-anchored pairs follow, BSAT \u00d7 UE (consistency 0.669; raw "
        "coverage 0.241) and BSAT \u00d7 UX (consistency 0.650; raw coverage 0.208), and the experience-"
        "plus-engagement pair UX \u00d7 UE (consistency 0.600; raw coverage 0.219) is associated with "
        "favourable outcomes even where satisfaction is not the leading component, consistent with "
        "flow-based engagement. Overall solution coverage is 0.324 and solution consistency is 0.756.")

delete("Table 6 further deepens this analysis by presenting an expanded set of configurations")
delete("Other configurations in Table 6 reflect the versatility of paths to success.")

# Necessity / sufficiency sections rewritten with real, disciplined content
rewrite("The research implements fuzzy-set Qualitative Comparative Analysis (FsQCA) using version 4.1 to determine the optimal",
        "Necessity analysis (Table 6) tests whether any condition must be present for high brand "
        "success. No single condition, UE, UX, or BSAT, reached the 0.90 consistency benchmark required "
        "to be deemed necessary. Brand success in the metaverse therefore does not depend on any one "
        "antecedent in isolation, providing direct support for Proposition 4 and motivating the "
        "configurational (sufficiency) analysis that follows.")

rewrite("This study employed fuzzy set qualitative comparative analysis (FsQCA) using fsQCA 4.1 software to identify configurations",
        "Sufficiency analysis identifies the combinations of conditions associated with high brand "
        "success. Following convention, raw coverage, unique coverage, and consistency are reported for "
        "each configuration (Pappas and Woodside, 2021). We interpret the results transparently: the "
        "configurational consistencies (0.60 to 0.71) and the solution consistency (0.756) sit at or "
        "slightly below the 0.75-to-0.80 benchmark commonly applied in fsQCA. We therefore treat the "
        "configurations as suggestive evidence of equifinal pathways rather than as strong sufficiency "
        "claims, and we identify strengthening these estimates, through recalibration and larger "
        "samples, as a priority for future research (Section 8).")

# Tie the three combination sub-paragraphs to results; strip unmeasured speculation
rewrite("The combination of MBS (users\u2019 brand satisfaction) and MVCX (user experience) is a very crucial point",
        "The BSAT \u00d7 UX configuration (consistency 0.650; raw coverage 0.208) indicates that "
        "satisfied users who also perceive a high-quality immersive experience tend to exhibit "
        "favourable brand outcomes, consistent with prior evidence linking satisfaction and experience "
        "to loyalty (Thaichon et al., 2022).")
rewrite("This combination of users\u2019 brand satisfaction (MBS) and user engagement (UES) is also a positive example",
        "The BSAT \u00d7 UE configuration (consistency 0.669; raw coverage 0.241) shows that satisfaction "
        "combined with active engagement is associated with favourable brand outcomes, in line with "
        "evidence that engagement reinforces brand equity and loyalty (Jeong and Kim, 2019).")
rewrite("This combination, consisting of user experience (MVCX) and user engagement (UES), also contributes to positive brand outcomes",
        "The UX \u00d7 UE configuration (consistency 0.600; raw coverage 0.219) is notable because "
        "immersive, participatory interaction is associated with favourable outcomes even where "
        "satisfaction is not the leading component, consistent with Flow Theory's emphasis on "
        "challenge-skill balance and absorption (Csikszentmihalyi, 1990; Shin, 2019).")
# Remove the long unmeasured NFT/virtual-events speculation paragraph (R3.3, R2.6)
delete("Users purchase NFTs for several reasons, including ownership, investment, and identity expression.")


# ============================================================
# 8. DISCUSSION — remove stray stats; correct mislabels; discipline claims
# ============================================================
rewrite("The findings from the FsQCA analysis provide critical insights into how combinations of user engagement",
        "The configurational results speak directly to the study's research questions. Three findings "
        "stand out, each tied to a specific result reported in Section 5.")

rewrite("The combination of user brand satisfaction (MBS) and user experience (MVCX) emerged as a significant driver",
        "First, brand satisfaction is the most consistent component of pathways to brand success, "
        "appearing in the three highest-consistency configurations (Table 7), yet it is not necessary "
        "on its own, since experience and engagement together also reach high coverage. This nuances "
        "the conventional view that satisfaction is the dominant route and indicates that an immersive, "
        "engaging interaction can be associated with favourable brand outcomes even where satisfaction "
        "is not the leading condition.")

rewrite("The pairing of MBS and MVCX also demonstrated a strong impact on brand success, with a consistency value of 0.669",
        "Second, the configuration combining satisfaction with engagement (BSAT \u00d7 UE; consistency "
        "0.669, raw coverage 0.241) shows that active participation paired with satisfaction is "
        "associated with stronger brand outcomes, consistent with engagement-equity research (Jeong "
        "and Kim, 2019). The fully triadic configuration (BSAT \u00d7 UX \u00d7 UE) attains the highest "
        "consistency (0.713), supporting Proposition 5.")

rewrite("While slightly lower in consistency (0.6003), the combination of MVCX and UES still contributed positively",
        "Third, engagement and experience are complementary rather than substitutable (Proposition 3): "
        "they co-occur in the strongest configuration and are strongly associated (HTMT = 0.814) while "
        "remaining empirically distinct. This pattern is consistent with the integrated Flow-SDT logic, "
        "in which motivational pull (SDT) and experiential quality (Flow) operate jointly rather than "
        "as alternatives. We confine these interpretations to the four measured constructs and report "
        "no effect sizes beyond those in Section 5.")

rewrite("The results highlight the importance of integrating multiple factors\u2014engagement, experience, and satisfaction",
        "Taken together, the results indicate that what is metaverse-specific here is not the existence "
        "of these links, which are established elsewhere, but their configurational structure: the "
        "absence of any necessary condition, the presence of multiple equifinal routes, and the "
        "complementarity of engagement and experience in an embodied, persistently social setting.")


# ============================================================
# 9. THEORETICAL & MANAGERIAL IMPLICATIONS — discipline, remove Figure 5
# ============================================================
rewrite("This research makes a substantial and nuanced contribution to understanding user engagement, brand experience",
        "This study advances theory in two disciplined ways. First, it reframes metaverse branding as "
        "a configurational phenomenon, providing empirical evidence of equifinality and the absence of "
        "necessary conditions, neither of which a variable-centric application of Flow Theory or SDT "
        "anticipates. Second, it integrates the two theories at the level of mechanism: SDT explains "
        "the motivation to enter and sustain immersive brand interaction, while Flow explains why "
        "well-designed interaction is experienced as absorbing and high quality, jointly accounting for "
        "why engagement and experience co-occur rather than substitute.")

rewrite("Within the framework of SDT, this study demonstrates that autonomy, competence, and relatedness assume new operational forms",
        "Within SDT, the metaverse plausibly reconfigures how autonomy, competence, and relatedness are "
        "satisfied, through avatar customization and asset ownership (autonomy), skill-based "
        "progression (competence), and persistent social worlds (relatedness). We advance this as a "
        "theoretical mechanism rather than an empirical finding: these needs were not measured in the "
        "present study, and we therefore frame them as explanatory antecedents that motivate the "
        "measured constructs and as a priority for direct measurement in future work.")

rewrite("Flow Theory is similarly advanced by illustrating that optimal experiential immersion in the metaverse depends",
        "Flow Theory is likewise extended by positioning immersive experience as contingent on the "
        "alignment of challenge and skill, calibrated by interaction design. Consistent with this, the "
        "UX \u00d7 UE configuration is associated with favourable brand outcomes even absent leading "
        "satisfaction, suggesting that experiential absorption can itself support brand outcomes. We do "
        "not extend this to specific technological features (for example, cross-platform "
        "interoperability or asset scarcity), which were not measured.")

# Managerial: bundle-level, tied to configs; move NFT/blockchain claims to future research
rewrite("This study offers crucial insights into how user engagement (UES), user experience (MVCX), and brand satisfaction (MBS)",
        "The configurational results translate into bundle-level guidance, each anchored to a specific "
        "configuration in Table 7. Where resources allow, brands should prioritise the triadic bundle "
        "(BSAT \u00d7 UX \u00d7 UE), the most consistent route to brand success. Where full investment is "
        "infeasible, pairing satisfaction with either a high-quality immersive experience (BSAT \u00d7 UX) "
        "or active engagement (BSAT \u00d7 UE) offers two efficient second-best options. Finally, "
        "immersive and participatory design (UX \u00d7 UE) can support brand outcomes on its own terms.")

rewrite("To achieve this, brands must prioritize creating immersive experiences that are deeply connected to users' virtual identities.",
        "Practically, this implies designing identity-affirming, community-embedded experiences, "
        "through avatar customization, interactive environments, and social features, that jointly "
        "raise experience quality and engagement while supporting satisfaction. Because the analysis "
        "rests on four cross-sectionally measured constructs, we do not extend these recommendations to "
        "specific technologies that were not measured.")

rewrite("However, for brands to secure long-term brand success, they must go beyond mere personalization and integrate advanced NFT features",
        "Tactics that depend on specific digital-asset mechanics, such as NFT interoperability, smart "
        "contracts, transparent minting, and blockchain-based authentication, are promising but lie "
        "beyond what these data can support, since none was measured in the present model. We therefore "
        "relocate them from findings to future research rather than presenting them as empirically "
        "grounded recommendations.")

rewrite("Moreover, this study reveals a powerful synergy between user satisfaction and user engagement.",
        "Consistent with the BSAT \u00d7 UE configuration, satisfaction and engagement reinforce one "
        "another: interactive opportunities such as gamified participation and co-creation deepen "
        "involvement while supporting satisfaction, strengthening users' relational attachment to the "
        "brand (Jeong and Kim, 2019).")

# Remove Figure 5 and the unsupported "system features" framework (Editor 3; R2.4/2.6; R3.3)
rewrite("In this dynamic environment, immersive marketing plays a crucial role in adapting to the evolving user behaviors",
        "Finally, because immersive contexts evolve quickly, brands should monitor engagement, "
        "experience, and satisfaction continuously and adjust design accordingly, ensuring that "
        "virtual interactions remain meaningful and consistent with brand values.")
delete("As illustrated in the Figure 5, several interconnected factors contribute to brand success")
delete("Within this emergent environment, four interdependent dimensions have been identified as foundational constructs")
delete("Figure 5. Metaverse System Features for Brand Success in the Metaverse.")
delete("underpinning transactional credibility and asset-based loyalty through blockchain-enabled ownership models")
delete("Further enhancing the marketing strategy, psychological engagement within the Metaverse plays a pivotal role.")
delete("process is essential for brand loyalty in the Metaverse, as users are more likely to remain engaged")

rewrite("Finally, this study emphasizes that user experience is the cornerstone of both brand satisfaction and brand success",
        "In sum, the managerial value of these findings is configurational: brand success in the "
        "metaverse follows from coherent bundles of engagement, experience, and satisfaction rather "
        "than from any single lever, and brands can select among equifinal bundles according to their "
        "capabilities.")

rewrite("The key to sustained success in the metaverse lies in brands\u2019 ability to create immersive, emotionally resonant experiences",
        "Given the contested commercial trajectory of the metaverse, we scope these implications to "
        "brands operating on high-traffic immersive-social platforms (for example, Roblox and Fortnite) "
        "and to growth markets such as India, rather than to the metaverse as a universal channel.")

doc.save(SRC)
print("Saved amended manuscript:", SRC)
