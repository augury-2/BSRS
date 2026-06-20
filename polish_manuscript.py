# -*- coding: utf-8 -*-
"""Final language/clarity polish pass (FT-50 readiness).

Rewrites paragraphs with run-ons, garbled prose, grammar errors, mis-citations
and grandiose filler into clear, concise academic English. All rewrites are
highlighted yellow; two redundant paragraphs are merged/removed. Clean sections
(abstract, literature review, tables, hypotheses, limitations) are left as-is.
"""
import docx
from docx.enum.text import WD_COLOR_INDEX

DOCX = "Manuscript Metaverse.docx"
d = docx.Document(DOCX)
paras = d.paragraphs  # stable object references


def set_segments(p, segments):
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    for text, hl in segments:
        if not text:
            continue
        run = p.add_run(text)
        if hl:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def rewrite(idx, expect, new_text):
    assert expect in paras[idx].text, f"para {idx} mismatch; expected {expect!r}"
    set_segments(paras[idx], [(new_text, True)])


def delete(idx, expect):
    assert expect in paras[idx].text, f"para {idx} delete mismatch; expected {expect!r}"
    el = paras[idx]._element
    el.getparent().remove(el)


# ---------------- INTRODUCTION ----------------
rewrite(21, "users interact with each other and virtual objects, creating unprecedented",
    "The metaverse is an immersive digital space in which users interact with one another and with "
    "virtual objects, creating new opportunities for brands to engage their audiences (Arya et al., 2024a; "
    "Zallio and Clarkson, 2022). Built on augmented reality (AR) and virtual reality (VR), it integrates "
    "digital information with the physical world and is reshaping how brands engage consumers, prompting a "
    "shift away from traditional marketing approaches (Arya et al., 2024a; Hennig-Thurau et al., 2023; "
    "Lee et al., 2022). By aligning these technologies with marketing objectives, brands can design "
    "digital experiences that deliver measurable value (Anderson et al., 2024).")

rewrite(22, "From a marketing perspective, managing these technologies",
    "From a marketing perspective, managing these technologies requires frameworks that ensure "
    "high-quality user experiences and integrate smoothly with existing digital platforms (Kopalle et al., "
    "2024). Unlike conventional digital channels, the metaverse enables multisensory brand experiences "
    "that blur physical and digital boundaries and link individual-level engagement to wider social "
    "systems. To realise its potential for consumer interaction and value creation, brands must adopt "
    "adaptive, forward-looking strategies that keep them competitive in a fast-changing environment "
    "(Marvi et al., 2024; Bilgihan et al., 2024; Hwang and Seo, 2025).")

rewrite(24, "explore new aesthetics in communication",
    "The metaverse allows brands to experiment with new forms of communication and to raise visibility "
    "through tailored, interactive experiences, while also requiring them to manage challenges such as "
    "data privacy and interoperability (Bilgihan et al., 2024; Koohang et al., 2023). This study addresses "
    "a gap in the literature by examining systematically how user engagement and immersive experience "
    "influence brand success in the metaverse. In these environments, consumer engagement is active and "
    "immersive rather than passive, changing how brands must design their strategies.")

rewrite(28, "transcend physical limitations, engaging global users",
    "The metaverse also lets brands transcend physical limitations and reach global audiences, with "
    "data-driven insights supporting consistent management of these interactions across regions "
    "(Dwivedi et al., 2023). By personalising experiences, brands can strengthen perceptions and "
    "satisfaction (Chandra et al., 2022); offerings tailored to individual needs are more likely to build "
    "emotional bonds and, in turn, word-of-mouth and advocacy.")

rewrite(31, "bridges theoretical frameworks like flow theory",
    "This study draws on Flow Theory and Self-Determination Theory (SDT) to extend their applicability to "
    "immersive digital ecosystems (Fan et al., 2022). Flow Theory (Csikszentmihalyi, 1990) has been "
    "applied to immersive contexts by Arya et al. (2024) to examine how a balance between user challenges "
    "and skills fosters deep engagement. SDT, in turn, explains users' intrinsic motivation in virtual "
    "environments (Deci and Ryan, 2000): satisfying the needs for autonomy (control over one's actions), "
    "competence (a sense of effectiveness) and relatedness (social connection) fosters intrinsic "
    "motivation and well-being (Ryan and Deci, 2017). In the metaverse, customisable avatars (autonomy), "
    "skill-based challenges (competence) and social spaces (relatedness) act as catalysts for deep "
    "engagement and, ultimately, brand success.")

# ---------------- METHODOLOGY ----------------
rewrite(75, "The methodology section of the research article explains",
    "This section sets out the study's methodology: the rationale for fuzzy-set qualitative comparative "
    "analysis (fsQCA), the identification and selection of variables, the data-collection procedure, the "
    "measures, the sample, and the analytical approach.")

rewrite(143, "a technique for coding known as",
    "To analyse the data on user engagement (UE), user experience (UX), brand satisfaction (BSAT) and "
    "brand success (BSUC), the raw responses were first calibrated, following Ragin (2008), into "
    "fuzzy-set membership scores that express the degree to which each case belongs to a given condition "
    "and to the outcome.")

rewrite(144, "the reactions were aligned involving three cut-off points",
    "The calibration followed established fsQCA practice (Douglas et al., 2020), converting the original "
    "Likert responses for UE, UX, BSAT and BSUC into fuzzy-set membership scores on a common 0\u20131 scale, "
    "as detailed below.")

rewrite(152, "Besides the limited potential to develop new methods",
    "As shown in Table 3, the intermediate solution models brand success as a function of brand "
    "satisfaction, user experience and user engagement, BSUC = f(BSAT, UX, UE). Raw coverage indicates "
    "the share of the outcome accounted for by each configuration: BSAT and UX together account for about "
    "20.84% of cases, BSAT and UE about 24.15%, and UX and UE about 21.90%. Unique coverage\u2014the share of "
    "cases uniquely explained by a configuration\u2014is 3.59% for BSAT \u00d7 UX, 6.90% for BSAT \u00d7 UE and 4.65% "
    "for UX \u00d7 UE.")

rewrite(153, "Consistency stands for the adherence to the model",
    "Consistency reflects how reliably a configuration is associated with the outcome; higher values "
    "indicate stronger sufficiency. As reported in Table 4, consistency is 0.650 for BSAT \u00d7 UX, 0.669 for "
    "BSAT \u00d7 UE and 0.600 for UX \u00d7 UE. For the intermediate solution as a whole, solution coverage is "
    "0.324\u2014indicating that the configurations jointly account for roughly 32% of the outcome\u2014and solution "
    "consistency is 0.756.")

rewrite(156, "We wanted to find out what it takes for a brand to succeed",
    "In the conceptual model, user experience, user engagement and brand satisfaction are the antecedent "
    "conditions for brand success. The subset\u2013superset analysis evaluates which combinations of these "
    "conditions are jointly sufficient for high brand success, rather than estimating a single directional "
    "path among them, thereby grounding the conceptual model in the empirical configurations.")

# ---------------- SUFFICIENCY NARRATIVE ----------------
rewrite(187, "is a very crucial point for the brand",
    "The combination of brand satisfaction (BSAT) and user experience (UX) is a key driver of strong "
    "brand outcomes, recording high raw coverage, unique coverage and consistency. Brand success is "
    "closely tied to customer fulfilment and the quality of the user's experience: satisfied users who "
    "enjoy well-designed brand experiences are more likely to remain loyal and to recommend the brand to "
    "others.")

rewrite(189, "is also a positive example for brand outcomes",
    "The combination of brand satisfaction (BSAT) and user engagement (UE) is likewise associated with "
    "favourable brand outcomes, with high raw coverage, unique coverage and consistency. Active "
    "participation and interaction with a brand are well-established drivers of positive outcomes: Jeong "
    "and Kim (2019) report a positive relationship between engagement and brand loyalty, and Ebrahim "
    "(2020) shows that social-media marketing can strengthen customer equity.")

rewrite(191, "also contributes to positive brand outcomes, although it shows slightly lower",
    "The combination of user experience (UX) and user engagement (UE) also contributes to favourable "
    "brand outcomes, though with slightly lower consistency than the other configurations. This pattern "
    "is consistent with prior work identifying the interplay of experience and engagement as central to "
    "brand success in immersive settings.")

# ---------------- IMPLICATIONS ----------------
rewrite(208, "makes a substantial and nuanced contribution",
    "This research contributes to understanding user engagement, experience and brand satisfaction in the "
    "metaverse by integrating two complementary frameworks: Self-Determination Theory (SDT) and Flow "
    "Theory. Situating these theories within immersive, digitally mediated environments\u2014characterised by "
    "persistent virtual communities, avatar-based identity and interactive brand interfaces\u2014clarifies how "
    "psychological needs and experiential states jointly shape engagement, satisfaction and brand "
    "outcomes. In contrast to earlier applications in traditional or gamified settings, the findings "
    "suggest that metaverse interactions involve a distinctive interplay between intrinsic motivation and "
    "experiential absorption, extending the reach of both theories into co-created virtual domains.")

rewrite(211, "The central theoretical innovation of this research",
    "The study's central theoretical contribution is the integration of SDT and Flow Theory into a single "
    "framework that links the motivational antecedents and the experiential dynamics of engagement in "
    "immersive environments. The configurational fsQCA results identify multiple equifinal pathways "
    "through which engagement, experiential quality and satisfaction combine to produce favourable brand "
    "outcomes, indicating that intrinsic-need fulfilment and challenge\u2013skill alignment operate as "
    "mutually reinforcing mechanisms.")

rewrite(215, "This study offers crucial insights into how user engagement",
    "The findings carry several implications for brand managers. User experience emerges as a key "
    "determinant of both brand satisfaction and favourable brand outcomes, with high-quality, immersive "
    "experiences appearing consistently in the configurations associated with brand success. A brand's "
    "ability to craft engaging, interactive and personalised experiences in the metaverse is therefore "
    "central to sustaining engagement and long-term brand success.")

rewrite(218, "Moreover, this study reveals a powerful synergy",
    "The results also point to a strong synergy between satisfaction and engagement: combining active "
    "participation with high satisfaction strengthens brand equity and loyalty (Jeong and Kim, 2019). For "
    "managers, this means creating interactive opportunities\u2014gamified interactions, co-creation and "
    "reward systems\u2014that deepen involvement while sustaining satisfaction. As users help shape the "
    "virtual environment, their sense of ownership and belonging grows, reinforcing emotional attachment "
    "and the relational capital that supports long-term loyalty.")

# Merge split paragraph 219 + 220 into one clean paragraph; delete 220
rewrite(219, "Further enhancing the marketing strategy, psychological engagement",
    "Psychological engagement is equally important. The metaverse allows brands to move beyond one-way "
    "communication and to use personalisation and avatar customisation to foster emotional resonance. "
    "Experiences aligned with consumers' self-concept and aspirations help users feel authentically "
    "represented within virtual spaces (Jung and Pawlowski, 2014), and the resulting emotional bond "
    "supports loyalty, as users tend to remain engaged with brands that reflect their real and "
    "aspirational identities.")

rewrite(221, "Finally, this study emphasizes that user experience is the cornerstone",
    "Finally, the study underscores that user experience is central to both brand satisfaction and brand "
    "success in the metaverse. To remain competitive, brands must ensure that their immersive experiences "
    "are emotionally resonant, consistent and aligned with users' values and identities; functional or "
    "visually appealing virtual assets are no longer sufficient. Brands that consistently deliver "
    "high-quality, interactive experiences are better placed to build long-term engagement and loyalty, "
    "reflecting a holistic approach in which engagement, experience and satisfaction reinforce one another "
    "to drive brand success.")

# ---------------- DELETIONS (redundant filler / merged) ----------------
delete(220, "process is essential for brand loyalty in the Metaverse")
delete(161, "On another level, we saw how those three variables")

d.save(DOCX)
print("Polish pass complete.")
