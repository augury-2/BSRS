#!/usr/bin/env python3
"""Fifth pass: implement the additive items from the revision strategy:
(1) explicit construct definitions, (2) a sharpened overarching research
question, (3) a note on marker-variable/CLF as complementary CMB checks."""
from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

SRC = "Manuscript Metaverse.docx"
doc = Document(SRC)


def find_p(sub):
    for p in doc.paragraphs:
        if sub in p.text:
            return p
    return None


def set_text(p, text):
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    p.add_run(text)


def insert_after(p, text, bold_lead=False):
    new_p = OxmlElement('w:p')
    p._p.addnext(new_p)
    np = Paragraph(new_p, p._parent)
    if text:
        np.add_run(text)
    return np


def inline(sub, old, new):
    p = find_p(sub)
    if p is None or old not in p.text:
        print("  [MISS inline]", sub[:45]); return
    set_text(p, p.text.replace(old, new))


# 1. Construct definitions, inserted at the head of the Theoretical Foundation section
tf = None
for p in doc.paragraphs:
    if p.text.strip() == 'Theoretical Foundation':
        tf = p; break
if tf is None:
    tf = find_p('Theoretical Foundation')
if tf is not None:
    defs = insert_after(tf,
        "We define the four focal constructs as follows. User engagement (UE) is the user's "
        "cognitive, emotional, and behavioural involvement with the brand during metaverse "
        "interaction. User experience (UX) is the perceived quality of that interaction, "
        "encompassing immersion and presence, navigability, sensory quality, and authenticity. "
        "Brand satisfaction (BSAT) is the user's overall evaluative judgement of the brand formed "
        "through metaverse encounters. Brand success (BSUC) is the set of brand-favourable outcomes "
        "these encounters produce, including continued use, choice and purchase intention, and "
        "advocacy or word of mouth. These definitions anchor the measurement model (Table 1) and the "
        "propositions developed below, and they hold the constructs conceptually distinct so that "
        "their relationships can be tested rather than assumed.")
    # add a small heading label paragraph before the definitions
    head = OxmlElement('w:p')
    defs._p.addprevious(head)
    Paragraph(head, defs._parent).add_run("3.1 Construct Definitions")
    print("inserted construct definitions")
else:
    print("  [WARN] Theoretical Foundation heading not found")

# Renumber the following theory subheadings to keep order (3.1 Flow -> 3.2, SDT label stays)
inline("Flow theory, as conceptualized by", "3.1 Flow theory", "3.2 Flow theory") if find_p("3.1 Flow theory") else None
fp = None
for p in doc.paragraphs:
    if p.text.strip() == '3.1 Flow theory':
        set_text(p, '3.2 Flow theory'); print("renumbered Flow heading"); break


# 2. Sharpen the overarching research question
inline("We address three research questions.",
       "We address three research questions.",
       "Our overarching question is: how do immersive, avatar-mediated metaverse interactions alter "
       "the configuration of user engagement, user experience, and brand satisfaction that leads to "
       "brand success? We decompose this into three research questions.")


# 3. Note marker-variable / common-latent-factor as complementary CMB checks
inline("provide convergent evidence that the constructs are empirically distinct rather than method artefacts",
       "provide convergent evidence that the constructs are empirically distinct rather than method artefacts.",
       "provide convergent evidence that the constructs are empirically distinct rather than method "
       "artefacts. Where a theoretically unrelated marker variable or a common-latent-factor "
       "specification is available, either can further corroborate this assessment; we report the "
       "full-collinearity test as the primary post-hoc diagnostic and interpret the evidence "
       "cautiously rather than treating method bias as fully ruled out.")


doc.save(SRC)
print("Saved:", SRC)
