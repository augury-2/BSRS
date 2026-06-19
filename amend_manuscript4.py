#!/usr/bin/env python3
"""Fourth pass: remove duplicated/low-quality leftover paragraphs and tighten
remaining verbose prose flagged by reviewers (R1, R3.12)."""
from docx import Document

SRC = "Manuscript Metaverse.docx"
doc = Document(SRC)


def find_p(sub):
    for p in doc.paragraphs:
        if sub in p.text:
            return p
    return None


def set_text(p, text):
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    p.add_run(text)


def rewrite(sub, new):
    p = find_p(sub)
    if p is None:
        print("  [MISS rewrite]", sub[:45]); return
    set_text(p, new)


def delete(sub):
    p = find_p(sub)
    if p is None:
        print("  [MISS delete]", sub[:45]); return
    p._element.getparent().remove(p._element)


def inline(sub, old, new):
    p = find_p(sub)
    if p is None or old not in p.text:
        print("  [MISS inline]", sub[:45]); return
    set_text(p, p.text.replace(old, new))


# 1. Remove duplicate research-question paragraph (keep the cleaner one at end of intro)
delete("Based on the above discussion, the following research questions are proposed")

# 2. Remove duplicated Cronbach paragraph in 5.2 (R1: alpha repeated)
delete("Reliability was assessed using Cronbach\u2019s \u03b1 and composite reliability. The Cronbach\u2019s \u03b1 values were 0.970")

# 3. Remove redundant/garbled calibration paragraphs in 5.3.1 (R3.12)
delete("Subsequently, the reactions were aligned involving three cut-off points")
delete("In this calibration process, the lower bound is set at 0")

# 4. Soften the qual/quant 'bridge' phrasing in 4.1 (R1, R2.8)
inline("Translating variables into fuzzy sets allows nuanced measurement",
       "bridging quantitative and qualitative dimensions",
       "capturing partial set membership and degree of presence")

# 5. Fix the garbled intermediate-solution paragraph (R3.12)
rewrite("This indicates the model used for the analysis, where BSAT, UX, and UE are the input variables",
        "Here BSAT, UX, and UE are the causal conditions and BSUC the outcome. Raw coverage indicates "
        "the share of the outcome accounted for by each combination: BSAT and UX jointly cover about "
        "20.8 percent, BSAT and UE about 24.1 percent, and UX and UE about 21.9 percent. Unique "
        "coverage, which isolates the outcome attributable to a combination net of the others, is 3.6 "
        "percent (BSAT \u00d7 UX), 6.9 percent (BSAT \u00d7 UE), and 4.7 percent (UX \u00d7 UE).")

# 6. Tighten circular subset-superset paragraphs
rewrite("The FsQCA subset-superset analysis was used to elucidate the link",
        "Subset-superset analysis was used to relate the four constructs to the conceptual model by "
        "empirically identifying the configurations of conditions associated with high brand success.")
rewrite("In the conceptual model, User Experience, User Engagement, and Brand Satisfaction are identified as the primary factors",
        "Within the model, user experience, user engagement, and brand satisfaction are the candidate "
        "conditions for brand success. The configurational analysis evaluates which combinations of "
        "these conditions are sufficient for the outcome, and whether any is necessary, thereby testing "
        "the propositions rather than assuming a single causal path.")
delete("On another level, we saw how those three variables interact and how certain scenarios involving them")

# 7. Fix the 3.1 Flow tail (remove 'structure of conducting the surveys'; honour scope condition)
rewrite("Contemporary metaverse research has established sophisticated connections",
        "Contemporary research connects flow to virtual brand experiences across several settings: "
        "e-commerce live streaming, where interactivity most strongly enhances flow (Wu et al., 2024); "
        "metaverse immersion in emerging markets (Liang et al., 2024); VR gaming, where deeper flow "
        "accelerates subjective time (Rutrecht et al., 2021); and short-video engagement (Lo and Lai, "
        "2023). Building on this work, we use flow as the experiential mechanism linking immersive "
        "interaction to engagement and experience quality. Consistent with our scope condition, flow's "
        "antecedents (challenge-skill balance, clear goals, immediate feedback) are treated as "
        "theoretical mechanisms rather than separately measured variables.")

# 8. Fix 4.2 triangulated-approach paragraph (remove 'qualitative input')
rewrite("Theoretical relevance and empirical evidence were central to variable selection.",
        "Variable selection prioritised theoretical relevance and demonstrated reliability; only "
        "constructs with strong conceptual grounding and established psychometric performance were "
        "retained, and a pilot study confirmed item clarity and reliability before full-scale "
        "administration. This sequence, literature review, item screening, and quantitative "
        "validation, ensured that the measures robustly captured user engagement, experience, "
        "satisfaction, and brand success.")

# 9. Tighten 6.1 'Unique insights' section (R1, R3.12)
rewrite("The analysis using FsQCA yields one-of-a-kind understandings",
        "fsQCA offers insights into the relationships among brand satisfaction (BSAT), user experience "
        "(UX), and user engagement (UE) that net-effects methods cannot, because its value lies in "
        "analysing combinations of conditions rather than the average contribution of each in isolation.")
rewrite("Combinations of Models: Factor wise Qualitative Comparative Analysis (QCA)",
        "Combinations rather than net effects. fsQCA identifies the joint conditions sufficient for high "
        "brand success, revealing equifinal pathways that variable-centric methods, focused on the "
        "average effect of each predictor, do not surface.")
rewrite("Consistency: The analysis delivers insights about consistency and raw coverage",
        "Consistency and coverage. Consistency indicates how reliably a configuration is associated with "
        "the outcome, while coverage indicates its empirical relevance; together they characterise both "
        "the dependability and the prevalence of each pathway.")
rewrite("Analysis: The FsQCA method enables a comparative analysis of the different combinations",
        "Comparative, configurational insight. By comparing configurations, the analysis shows how "
        "conditions combine to produce brand success, giving managers actionable, bundle-level guidance "
        "rather than a single ranked predictor.")

# 10. Discipline the two overblown 7.1 paragraphs
rewrite("The central theoretical innovation of this research lies in the integration of SDT and Flow Theory into a dual-theoretical",
        "The central theoretical contribution is the integration of SDT and Flow Theory into a single "
        "configurational account of metaverse branding. Using fsQCA, the study identifies multiple "
        "equifinal pathways in which engagement, experiential quality, and satisfaction combine to "
        "produce high brand success, and shows that no single condition is necessary. This positions "
        "intrinsic motivation (SDT) and experiential absorption (Flow) as complementary, jointly "
        "operating mechanisms rather than competing explanations.")
rewrite("Collectively, these contributions advance both marketing and psychological theory by demonstrating",
        "In doing so, the study extends both theories into technologically mediated, user-driven "
        "domains while remaining disciplined about scope: the configurational evidence concerns four "
        "measured constructs, and the SDT and Flow needs and antecedents are advanced as explanatory "
        "framing for future, direct empirical testing.")

doc.save(SRC)
print("Saved:", SRC)
