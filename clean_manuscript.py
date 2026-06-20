# -*- coding: utf-8 -*-
"""Surgically clean the metaverse-branding manuscript for FT-50 submission.

- Removes unmeasured-construct claims (NFTs/blockchain/smart-contracts/ROI) shown as findings
- Removes untraceable statistics (71%, 73%, mean 4.48/3.89, r=0.7) and orphan participant quotes
- Fixes contradictory Necessity/Sufficiency subsections to match real fsQCA results
- Removes Figure 5 (system-features) + its blockchain/NFT discussion and orphan fragment
- Removes redundant fsQCA-tutorial filler and grandiose passages
- Removes phantom figure references (non-existent "Figure 2"/"Figure 3" callouts)
- All rewritten/added text is highlighted yellow; pure deletions are removed.
"""
import docx
from docx.enum.text import WD_COLOR_INDEX

DOCX = "Manuscript Metaverse.docx"
d = docx.Document(DOCX)
paras = d.paragraphs  # snapshot list of paragraph objects (stable references)


def set_segments(p, segments):
    """Replace a paragraph's content with given (text, highlight) segments."""
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    for text, hl in segments:
        if text == "":
            continue
        run = p.add_run(text)
        if hl:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def rewrite(idx, new_text):
    set_segments(paras[idx], [(new_text, True)])


def truncate_at(idx, marker, replacement=None):
    """Keep text before `marker`; optionally append a highlighted replacement clause."""
    full = paras[idx].text
    pos = full.find(marker)
    assert pos != -1, f"marker not found in para {idx}: {marker!r}"
    prefix = full[:pos].rstrip()
    segs = [(prefix, False)]
    if replacement:
        segs = [(prefix + " ", False), (replacement, True)]
    set_segments(paras[idx], segs)


def replace_phrase(idx, old, new):
    full = paras[idx].text
    pos = full.find(old)
    assert pos != -1, f"phrase not found in para {idx}: {old!r}"
    before = full[:pos]
    after = full[pos + len(old):]
    set_segments(paras[idx], [(before, False), (new, True), (after, False)])


def delete(idx):
    el = paras[idx]._element
    el.getparent().remove(el)


# ---------------- REWRITES (full paragraph, highlighted) ----------------

rewrite(32,
    "This study advances a configurational account of brand success in the metaverse in which "
    "user engagement (UE) and user experience (UX) act as antecedents of brand satisfaction (BSAT), "
    "which in turn is associated with brand success (BSUC). Rather than treating these as isolated "
    "drivers, the analysis examines how they combine, consistent with the view that brand outcomes in "
    "immersive, avatar-mediated environments emerge from configurations of engagement, experience and "
    "satisfaction rather than from any single factor.")

rewrite(83,
    "By systematically applying fsQCA, the study identifies the engagement\u2013experience\u2013satisfaction "
    "configurations associated with brand success in virtual environments, complementing the "
    "measurement model with a configurational account and offering practical guidance for managers "
    "designing resilient metaverse strategies.")

rewrite(95,
    "The Sandbox focuses on voxel-based content creation, allowing users to design, own and monetise "
    "experiences; major brands such as Adidas have used the platform for interactive campaigns "
    "(Brandes and D\u00f6lp, 2025). Across these platforms, the study considered four primary types of "
    "metaverse activity: trading digital assets to enhance virtual identities or for investment "
    "purposes (Lee and Shen, 2024); attending virtual events; interacting with other users through "
    "avatars to build communities or collaborate on shared goals (Hennig-Thurau et al., 2023); and "
    "designing digital assets or developing interactive experiences that contribute to the virtual "
    "ecosystem (Bilgihan et al., 2024).")

rewrite(181,
    "The necessity analysis examines whether any single condition must be present for high brand "
    "success. Following standard practice, a condition is treated as necessary only when its "
    "consistency exceeds 0.90 (Schneider and Wagemann, 2012). In the present data no individual "
    "condition met this threshold: user engagement, user experience and brand satisfaction recorded "
    "single-condition consistency well below the cut-off (0.373, 0.362 and 0.410, respectively). High "
    "brand success is therefore not attributable to any one necessary condition but arises from "
    "combinations of conditions, motivating the sufficiency analysis that follows.")

rewrite(184,
    "The sufficiency analysis identifies the configurations of conditions that are jointly sufficient "
    "for high brand success. The three calibrated conditions\u2014user engagement (UE), user experience "
    "(UX) and brand satisfaction (BSAT)\u2014were entered into the truth table, with consistency and "
    "frequency thresholds applied in line with established guidance (Ragin, 2008; Schneider and "
    "Wagemann, 2012). The resulting configurations and their consistency and coverage values are "
    "reported in Tables 5 and 6: brand satisfaction is present in every high-consistency sufficient "
    "path, while engagement and experience contribute in combination rather than in isolation.")

rewrite(192,
    "User engagement in these settings is multifaceted, spanning asset trading, attendance at virtual "
    "events, social interaction and content creation. These behaviours reflect a combination of "
    "psychological, social and experiential motivations and are captured in the present study through "
    "the user-engagement and user-experience constructs rather than as separate platform-specific "
    "variables.")

rewrite(199,
    "The combination of brand satisfaction (BSAT) and user experience (UX) emerged as a significant "
    "configuration for brand success, with a consistency value of 0.650 and raw coverage of 20.84%. "
    "This indicates that satisfied users who also perceive high-quality, immersive experiences are "
    "more likely to develop loyalty and advocate for the brand. The pattern is consistent with prior "
    "work showing that customer satisfaction and positive user experience foster emotional connection "
    "and longer-term engagement (Thaichon et al., 2022; Bilgihan et al., 2024). Brands should therefore "
    "prioritise seamless, immersive virtual experiences that align with user preferences in order to "
    "strengthen satisfaction and engagement. At the same time, immersive environments raise ethical "
    "considerations\u2014including data-privacy risks and the potential for excessive use\u2014that brands and "
    "platform designers should address.")

rewrite(201,
    "The combination of user experience (UX) and user engagement (UE), while slightly lower in "
    "consistency (0.600), also contributed positively to brand outcomes, with raw coverage of 21.90%. "
    "This configuration indicates that engagement is most effective when paired with high-quality "
    "experience, even where satisfaction is not the dominant condition\u2014consistent with Flow Theory's "
    "emphasis on deep engagement through challenge\u2013skill balance (Csikszentmihalyi, 1990; Shin, 2019). "
    "To leverage this pathway, brands should design engaging virtual spaces that captivate users "
    "through interactive, well-paced experiences.")

rewrite(221,
    "To achieve this, brands should create immersive experiences that connect to users' virtual "
    "identities. Personalisation is central: features such as customisable avatars and interactive "
    "virtual environments allow users to express identity within the metaverse and strengthen their "
    "sense of presence, which in turn reinforces the emotional bond between user and brand. By aligning "
    "a brand's virtual presence with users' self-concept, brands can cultivate connection that supports "
    "satisfaction. When users perceive their virtual experiences as tailored to them, they are more "
    "likely to engage deeply, exhibit loyalty and advocate for the brand within their communities "
    "(Thaichon et al., 2022; Xie and Muralidharan, 2023).")

rewrite(234,
    "Sustained success in the metaverse rests on brands' ability to create immersive, emotionally "
    "resonant experiences that integrate personalisation, interactive marketing and trust-building. "
    "Brands that master these dimensions will be well positioned to foster deep user engagement, "
    "long-term loyalty and a competitive edge in this rapidly evolving virtual space.")

# ---------------- PHRASE / PHANTOM-REFERENCE EDITS ----------------

# Para 52: remove phantom "Figure 2 illustrates these dynamics..." callout, keep substance (reworded)
truncate_at(52, "Figure 2 illustrates these dynamics",
            "Interactivity\u2014through avatars, two-way communication and community identity\u2014further "
            "enhances belonging and loyalty (Hennig-Thurau et al., 2023; Payal et al., 2024).")

# Para 58: remove phantom "Figure 3 illustrates these moderating effects..." callout (pure removal)
truncate_at(58, "Figure 3 illustrates these moderating effects")

# Para 213: drop NFTs from the contextual description in theoretical contributions
replace_phrase(213,
    "virtual communities, non-fungible tokens (NFTs), and interactive brand interfaces",
    "persistent virtual communities, avatar-based identity, and interactive brand interfaces")

# ---------------- DELETIONS (irrelevant / redundant / unsupported) ----------------
# 6.1 fsQCA-tutorial filler (heading + generic explanation paragraphs)
for i in [205, 206, 207, 208, 209]:
    delete(i)
# Figure 5 block: KPI filler, system-features intro, four-dimensions text, image, caption, orphan fragment
for i in [224, 225, 226, 227, 228, 229, 230]:
    delete(i)

d.save(DOCX)
print("Saved cleaned manuscript.")
