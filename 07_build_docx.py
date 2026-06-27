from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ---------- base styles ----------
normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15

def set_cell_bg(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0x1F, 0x2D, 0x3D)
    return h

def add_para(text, italic=False, size=11, bold=False, align=None, space_after=6):
    p = doc.add_paragraph()
    if align: p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.italic = italic; run.bold = bold; run.font.size = Pt(size)
    return p

def add_table(headers, rows, note=None, caption=None, col_bold_first=False):
    if caption:
        cap = doc.add_paragraph()
        cap.paragraph_format.space_before = Pt(8); cap.paragraph_format.space_after = Pt(2)
        r = cap.add_run(caption); r.bold = True; r.font.size = Pt(10.5)
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True; run.font.size = Pt(9.5); run.font.name = "Times New Roman"
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(hdr[i], "D9E2F3")
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(9.5); run.font.name = "Times New Roman"
            if col_bold_first and i == 0: run.bold = True
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if i>0 else WD_ALIGN_PARAGRAPH.LEFT
    if note:
        np_ = doc.add_paragraph()
        np_.paragraph_format.space_before = Pt(2)
        r = np_.add_run("Note. " + note); r.italic = True; r.font.size = Pt(9)
    return t

# ============ TITLE ============
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title.add_run("Results")
tr.bold = True; tr.font.size = Pt(16); tr.font.name = "Times New Roman"
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub.add_run("A Self-Determination-Theory Model of Brand Outcomes in Avatar-Mediated Environments: "
                 "PLS-SEM and fsQCA Evidence (N = 312)")
sr.italic = True; sr.font.size = Pt(11)
doc.add_paragraph()

add_para("This section reports the empirical evaluation of the Self-Determination-Theory (SDT) based model "
         "of avatar-mediated brand outcomes. We first describe the sample and basic descriptive statistics, "
         "then assess the reflective measurement model, test the structural hypotheses (H1a-H4) using partial "
         "least squares structural equation modeling (PLS-SEM) with bootstrapping, and finally examine the "
         "configurational proposition (P1) using fuzzy-set qualitative comparative analysis (fsQCA). Tables and "
         "the figure are referenced in their order of appearance.")

# Transparency note box
tn = doc.add_paragraph()
tn.paragraph_format.left_indent = Inches(0.2)
r = tn.add_run("Transparency note (read first). ")
r.bold = True; r.font.size = Pt(10.5)
r2 = tn.add_run("The hypotheses below were stated a priori from SDT and were revised only where the data forced a "
    "change. The measurement model is exemplary, but the data do not support any of the hypothesized structural "
    "or configurational relationships: the four latent constructs are essentially statistically orthogonal in "
    "this sample. We report this outcome fully, justify every threshold, and discuss the implications for both "
    "the model and the dataset. No cases were altered, trimmed, or re-weighted to manufacture effects.")
r2.font.size = Pt(10.5)

# ============ 1. SAMPLE ============
add_heading("1. Sample characteristics and descriptive statistics", 1)
add_para("The analysis is based on N = 312 complete responses collected with reflective, 7-point Likert-type "
    "items (1 = strongly disagree to 7 = strongly agree). The dataset contained no missing values across the 18 "
    "focal indicators and the two attitudinal control items (ATT_1, ATT_2), so no imputation was required. "
    "Univariate screening showed all items spanning the full 1-7 range with means clustered near the scale "
    "midpoint (item M range = 3.85-4.15) and substantial dispersion (item SD range = 1.83-2.01). Item-level "
    "skewness (|skew| <= 0.14) and kurtosis (range -1.22 to -1.02) were within accepted limits for symmetric, "
    "slightly platykurtic distributions. A multivariate outlier screen using Mahalanobis distance over the 18 "
    "indicators flagged no cases at the conservative p < .001 criterion (max D-squared = 39.54 < chi-square(18) "
    "= 42.31). Because PLS-SEM makes no multivariate-normality assumption and bootstrapping is used for "
    "inference, the modest non-normality is inconsequential.")
add_para("Construct composites were formed as item means: UE = mean(UE1-UE5), UX = mean(UX1-UX5), "
    "BSAT = mean(BSAT1-BSAT4), BSUC = mean(BSUC1-BSUC4). Composite descriptive statistics and their zero-order "
    "correlations (including the attitudinal controls) are reported in Table 1.")

add_table(
    ["Variable", "M", "SD", "1", "2", "3", "4", "5", "6"],
    [["1. UE","3.99","1.71","1.00","","","","",""],
     ["2. UX","3.96","1.65",".01","1.00","","","",""],
     ["3. BSAT","4.05","1.69","-.04",".08","1.00","","",""],
     ["4. BSUC","4.13","1.63",".04","-.05",".01","1.00","",""],
     ["5. ATT_1","3.85","1.84","-.12",".02","-.04","-.06","1.00",""],
     ["6. ATT_2","3.92","1.87","-.16",".07","-.09","-.01",".64","1.00"]],
    caption="Table 1. Means, standard deviations, and correlations among composites (N = 312).",
    note="No correlation among the four focal constructs reaches |r| = .09; none is significant at alpha = .05 "
         "(critical |r| ~ .111 for n = 312). The only sizeable association is between the two attitudinal control "
         "items (r = .64).",
    col_bold_first=True)
add_para("The descriptive picture already foreshadows the central finding: although each construct is "
    "well-measured (Section 2), the constructs are uncorrelated with one another.")

# ============ 2. MEASUREMENT ============
add_heading("2. Measurement model (PLS-SEM)", 1)
add_para("The model was estimated with PLS-SEM (Mode A, reflective measurement; path-weighting inner scheme), "
    "appropriate for a predictive, composite-based evaluation of an SDT nomological network. All four constructs "
    "(UE, UX, BSAT, BSUC) were specified as reflective, consistent with conceptualizing UE and UX as "
    "interaction-level manifestations of SDT need satisfaction, BSAT as the evaluative outcome, and BSUC as the "
    "brand-level outcome.")

add_heading("2.1 Indicator reliability and internal consistency", 2)
add_para("All standardized outer loadings substantially exceed the .70 benchmark (range = .831-.946); no item "
    "fell into the .40-.70 consider-removal band or below, so every indicator was retained. Internal consistency "
    "is excellent and uniform: Cronbach's alpha = .902-.934 and composite reliability (CR/rho-c) = .929-.942, all "
    "above .70 and below the .95 redundancy ceiling. Table 2 reports indicator properties with alpha, CR, and "
    "average variance extracted (AVE).")

add_table(
    ["Construct","Item","Loading","Cronbach's alpha","CR (rho-c)","AVE"],
    [["User Engagement (UE)","UE1",".832",".934",".942",".764"],
     ["","UE2",".862","","",""],["","UE3",".946","","",""],
     ["","UE4",".881","","",""],["","UE5",".845","","",""],
     ["User Experience (UX)","UX1",".860",".925",".943",".766"],
     ["","UX2",".900","","",""],["","UX3",".889","","",""],
     ["","UX4",".854","","",""],["","UX5",".874","","",""],
     ["Brand Satisfaction (BSAT)","BSAT1",".901",".913",".936",".786"],
     ["","BSAT2",".877","","",""],["","BSAT3",".857","","",""],
     ["","BSAT4",".910","","",""],
     ["Brand Success (BSUC)","BSUC1",".904",".902",".929",".767"],
     ["","BSUC2",".895","","",""],["","BSUC3",".871","","",""],
     ["","BSUC4",".831","","",""]],
    caption="Table 2. Measurement model: loadings, reliability, and AVE.",
    note="All loadings significant at p < .001 (bootstrap, 5,000 subsamples).")

add_heading("2.2 Convergent reliability (AVE)", 2)
add_para("AVE ranges from .764 to .786, exceeding the .50 criterion for every construct. Consistent with current "
    "measurement guidance, we interpret AVE as a reliability index - the proportion of indicator variance "
    "captured by the construct relative to measurement error (here ~76-79%) - rather than as a stand-alone test "
    "of validity.")

add_heading("2.3 Discriminant validity", 2)
add_para("Discriminant validity was assessed with the heterotrait-monotrait ratio of correlations (HTMT; Table "
    "3). All HTMT values are extraordinarily low (range = .032-.094), far below both the conservative .85 and "
    "liberal .90 thresholds. The Fornell-Larcker criterion is likewise satisfied: each construct's square-root "
    "AVE (~.87-.89) vastly exceeds its correlations with the other constructs (|r| <= .10). Discriminant validity "
    "is therefore not merely adequate but extreme - a direct consequence of the constructs being empirically "
    "independent.")
add_table(
    ["", "UE", "UX", "BSAT", "BSUC"],
    [["UE","-","","",""],
     ["UX",".041","-","",""],
     ["BSAT",".050",".094","-",""],
     ["BSUC",".042",".073",".032","-"]],
    caption="Table 3. Heterotrait-monotrait ratios (HTMT).",
    note="All HTMT < .85 (and < .90). The near-zero values indicate the constructs share virtually no common variance.",
    col_bold_first=True)

add_heading("2.4 Common method bias", 2)
add_para("Common method bias (CMB) was evaluated with Kock's full-collinearity approach: each construct was "
    "regressed on all others and the variance inflation factor (VIF) inspected. All full-collinearity VIFs are "
    "~1.00 (UE = 1.006, UX = 1.014, BSAT = 1.014, BSUC = 1.008), far below the 3.3 threshold (Table 4). There is "
    "thus no evidence of common method bias or pathological collinearity.")
add_table(
    ["Construct","VIF"],
    [["UE","1.006"],["UX","1.014"],["BSAT","1.014"],["BSUC","1.008"]],
    caption="Table 4. Full-collinearity VIFs (CMB check).",
    note="All VIF < 3.3; no common method bias indicated.",
    col_bold_first=True)
add_para("Measurement-model summary. The reflective measurement model is, by every standard criterion, "
    "exemplary: high and significant loadings, alpha and CR well above .70, AVE near .77, HTMT below .10, and no "
    "CMB. The instrument is reliable and the constructs are sharply distinct. This clean measurement foundation "
    "makes the structural results that follow interpretable rather than an artifact of poor measurement.",
    bold=False)

# ============ 3. STRUCTURAL ============
add_heading("3. Structural model and hypothesis tests (PLS-SEM)", 1)
add_para("The structural model specified UE -> BSAT, UX -> BSAT, BSAT -> BSUC, UE -> BSUC, and UX -> BSUC. The "
    "two attitudinal items (ATT_1, ATT_2) were entered as covariates on BSUC. Significance was assessed with "
    "bootstrapping using 5,000 subsamples (percentile 95% confidence intervals; two-tailed). Sign indeterminacy "
    "across bootstrap runs was resolved by aligning each construct's outer-weight signs to the full-sample "
    "solution.")

add_heading("3.1 Explained variance", 2)
add_para("The endogenous constructs are essentially unexplained: R-squared(BSAT) = .013 (adj. = .006) and "
    "R-squared(BSUC) = .008 (adj. = -.002). Adding the controls left R-squared(BSUC) unchanged at .013. By "
    "Cohen's benchmarks these values are below even the small threshold; the antecedents account for roughly 1% "
    "of the variance in the outcomes. Predictive relevance was therefore not pursued, as there is no systematic "
    "variance to predict.")

add_heading("3.2 Hypothesis tests", 2)
add_para("Table 5 reports the direct and indirect (mediation) effects. Figure 1 depicts the structural model "
    "with standardized coefficients and R-squared values.")
add_table(
    ["Hypothesis","Relationship","beta","t","p","95% CI","Supported?"],
    [["H1a","UE -> BSAT","-.045","0.58",".561","[-.166, .112]","No"],
     ["H1b","UX -> BSAT",".103","1.37",".171","[-.096, .219]","No"],
     ["H2","BSAT -> BSUC",".029","0.44",".663","[-.104, .143]","No"],
     ["H3a","UE -> BSUC",".062","0.79",".428","[-.111, .173]","No"],
     ["H3b","UX -> BSUC","-.060","0.76",".448","[-.185, .120]","No"],
     ["H4a","UE -> BSAT -> BSUC (indirect)","-.001","0.22",".827","[-.013, .012]","No"],
     ["H4b","UX -> BSAT -> BSUC (indirect)",".003","0.35",".724","[-.013, .020]","No"]],
    caption="Table 5. Structural model results and hypothesis tests (N = 312; 5,000 bootstrap subsamples).",
    note="Controls on BSUC were also non-significant: ATT_1 (beta = -.091, t = 1.23, p = .221) and ATT_2 "
         "(beta = .067, t = 0.89, p = .375).",
    col_bold_first=True)

# Figure 1
figcap = doc.add_paragraph(); figcap.paragraph_format.space_before = Pt(8)
fr = figcap.add_run("Figure 1. PLS-SEM structural model with standardized path coefficients and R-squared.")
fr.bold = True; fr.font.size = Pt(10.5)
pic = doc.add_paragraph(); pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
pic.add_run().add_picture("Figure1_structural_model.png", width=Inches(6.2))
add_para("All five direct paths are non-significant, and both endogenous R-squared values are near zero.",
         italic=True, size=9)

add_heading("3.3 Interpretation and hypothesis revision", 2)
add_para("Every hypothesized relationship is rejected. H1a/H1b: neither user engagement nor user experience "
    "predicts brand satisfaction; both 95% CIs include zero (UX -> BSAT is the strongest path at beta = .103 but "
    "still p = .171). H2: brand satisfaction does not predict brand success (beta = .029, p = .663). H3a/H3b: no "
    "direct engagement- or experience-to-success effect emerges. H4a/H4b: mediation requires non-trivial a- and "
    "b-paths; because both are null, the indirect effects are effectively zero (|beta| <= .003) with CIs spanning "
    "zero - there is no mediation, neither full nor partial.")
add_para("Following the pre-registered revise-only-if-forced rule, the data force the strongest possible "
    "revision: rather than dropping a single non-significant path and retaining a parsimonious model, the entire "
    "SDT-based path structure fails to replicate in this sample. We therefore do not reframe individual paths "
    "(e.g., demoting H3a/H3b to a satisfaction-only chain), because the mediator pathway (H1 and H2) is itself "
    "null. The honest conclusion is that, in these data, UE, UX, BSAT, and BSUC behave as mutually independent "
    "constructs.")

# ============ 4. fsQCA ============
add_heading("4. Configurational analysis (fsQCA)", 1)
add_para("To test the equifinality proposition P1, the four composites were calibrated into fuzzy sets using the "
    "direct method with three anchors appropriate to the 7-point scale: full membership = 6.0, crossover = 4.0, "
    "full non-membership = 2.0 (log-odds anchored at 0.95/0.50/0.05). The resulting fuzzy scores span the full "
    "[0,1] interval, with crossover near the empirical median (calibrated means ~.50 for UE/UX, .51 for BSAT, "
    ".53 for BSUC), confirming well-behaved calibration. The outcome was high brand success (BSUC_f).")

add_heading("4.1 Analysis of necessary conditions", 2)
add_para("A condition is necessary if its consistency for the outcome is >= .90. As shown in Table 6, no "
    "condition - present or absent - approaches necessity (all consistencies between .56 and .60). High BSUC can "
    "occur with or without high UE, UX, or BSAT; no single SDT construct is a prerequisite for brand success.")
add_table(
    ["Condition","Consistency","Coverage"],
    [["UE",".580",".614"],["~UE",".575",".607"],
     ["UX",".559",".597"],["~UX",".603",".631"],
     ["BSAT",".584",".605"],["~BSAT",".565",".610"]],
    caption="Table 6. Analysis of necessary conditions for high BSUC.",
    note="'~' denotes negation (low membership). Necessity benchmark = .90. No condition qualifies.",
    col_bold_first=True)

add_heading("4.2 Analysis of sufficient configurations", 2)
add_para("A truth table was constructed for the three causal conditions UE_f, UX_f, BSAT_f (the outcome BSUC_f "
    "is not a causal condition), yielding 2^3 = 8 configurations. Cases were assigned to the corner in which "
    "their membership exceeded 0.5. We applied a frequency threshold of 3 and a raw-consistency threshold of .80, "
    "the conventional minima for a sample of this size. Table 7 reports the full truth table.")
add_table(
    ["UE","UX","BSAT","n (cases)","Raw consistency"],
    [["present","present","present","42",".711"],
     ["absent","absent","absent","39",".739"],
     ["absent","present","present","38",".699"],
     ["present","absent","present","34",".762"],
     ["absent","absent","present","32",".713"],
     ["absent","present","absent","30",".714"],
     ["present","present","absent","29",".688"],
     ["present","absent","absent","28",".723"]],
    caption="Table 7. Truth table for UE, UX, BSAT -> high BSUC.",
    note="All eight configurations clear the frequency threshold (n >= 28), but none reaches the .80 consistency "
         "threshold (observed range = .688-.762). The truth table yields no rows coded 1, and Quine-McCluskey "
         "minimization returns an empty solution (no sufficient term).")
add_para("Because the standard procedure produces no solution, we additionally inspected the sufficiency of "
    "single and combined presence conditions directly (Table 8). Consistency rises monotonically with the number "
    "of co-present high conditions - from ~.60 for single conditions to .711 for the fully triadic recipe "
    "(UE*UX*BSAT) - exactly the qualitative pattern P1 anticipates, but the ceiling (.711) remains below any "
    "defensible sufficiency cut-off, and coverage of the triadic term is low (.262). For completeness, no "
    "configuration is sufficient for the negated outcome either (all consistencies <= .65).")
add_table(
    ["Configuration","Consistency","Raw coverage"],
    [["UE",".614",".580"],["UX",".597",".559"],["BSAT",".605",".584"],
     ["UE * UX",".653",".360"],["UE * BSAT",".683",".381"],
     ["UX * BSAT",".653",".382"],["UE * UX * BSAT",".711",".262"]],
    caption="Table 8. Sufficiency of selected configurations for high BSUC.",
    col_bold_first=True)

add_heading("4.3 Interpretation relative to P1", 2)
add_para("P1 holds that multiple configurations of high UE, high UX, and high BSAT are sufficient for high BSUC, "
    "with no single condition necessary. The data deliver half of this proposition and reject the other half. "
    "No single condition is necessary - consistent with the no-necessity clause of P1 (Table 6). However, no "
    "configuration is sufficient: not one of the eight corners - including the theoretically privileged triadic "
    "recipe - meets the .80 consistency standard (Table 7). There are therefore no equifinal sufficient paths to "
    "brand success.")
add_para("Consequently, P1 is not supported. We deliberately refrain from reporting C1/C2/C3 solution "
    "configurations, because doing so would require lowering the consistency threshold below accepted standards "
    "(no corner reaches even .77) and would misrepresent noise as an equifinal causal structure. The "
    "configurational evidence converges with the PLS-SEM evidence: brand success is, in this sample, unrelated "
    "to the SDT antecedents whether modeled correlationally or set-theoretically.")

# ============ 5. SUMMARY ============
add_heading("5. Summary of findings", 1)
add_table(
    ["#","Proposition / Hypothesis","Evidence","Verdict"],
    [["-","Measurement quality (reliability, convergent reliability, discriminant validity, CMB)",
      "alpha/CR > .90, AVE > .76, HTMT < .10, VIF ~ 1.0","Strongly met"],
     ["H1a","UE -> BSAT","beta = -.045, p = .561","Not supported"],
     ["H1b","UX -> BSAT","beta = .103, p = .171","Not supported"],
     ["H2","BSAT -> BSUC","beta = .029, p = .663","Not supported"],
     ["H3a","UE -> BSUC","beta = .062, p = .428","Not supported"],
     ["H3b","UX -> BSUC","beta = -.060, p = .448","Not supported"],
     ["H4","BSAT mediates UE/UX -> BSUC","indirect betas ~ 0, CIs include 0","Not supported"],
     ["P1","Equifinal sufficient configs; no necessary condition","No necessity AND no sufficiency","Not supported"]],
    caption="Table 9. Summary of hypothesis and proposition tests.",
    col_bold_first=True)

add_para("Overall. The MTVS data yield a paradoxical and noteworthy pattern: each construct is measured with "
    "near-textbook reliability and validity, yet the constructs are mutually orthogonal. Within-construct item "
    "correlations average ~.70-.74, whereas between-construct correlations average ~.02-.07 (max |r| = .10). As a "
    "result, every SDT-derived structural hypothesis (H1a-H4) and the configurational proposition (P1) are "
    "disconfirmed: there is no direct effect, no mediation, and no sufficient configuration linking user "
    "engagement, user experience, or brand satisfaction to brand success.")
add_para("Implication and caution. Two readings are possible, and both should be stated transparently. "
    "(1) Substantive: in this avatar-mediated setting, interaction-level need satisfaction (UE, UX) and "
    "evaluative satisfaction (BSAT) may be decoupled from brand-level success (BSUC), challenging the assumed SDT "
    "causal chain. (2) Data-integrity: the simultaneous combination of high internal consistency with near-zero "
    "inter-construct correlations is statistically unusual for genuinely related theoretical constructs and is "
    "the signature one would expect if the four blocks were generated or sampled independently. We flag this "
    "pattern explicitly. Before advancing substantive SDT claims, we recommend verifying the data-collection and "
    "data-assembly pipeline (e.g., possible block-wise shuffling or independent simulation of construct blocks) "
    "and, if confirmed valid, replicating on an independent sample. As reported, the responsible conclusion is "
    "that these data do not support the hypothesized SDT model of avatar-mediated brand success.")

# Methods note
add_heading("Methods note (reproducibility)", 2)
add_para("PLS-SEM was estimated with a Mode-A, path-weighting implementation (converging in < 20 iterations); "
    "loadings are indicator-construct correlations, rho-c and AVE follow the standard reflective formulas, HTMT "
    "follows Henseler et al. (2015), and full-collinearity VIFs follow Kock (2015). Inference used 5,000 "
    "bootstrap subsamples with construct-wise sign correction (seed = 20240627). fsQCA used direct calibration "
    "(anchors 6/4/2), Ragin-style fuzzy consistency/coverage, a corner-membership truth table (frequency >= 3, "
    "consistency >= .80), and Quine-McCluskey minimization. All scripts and intermediate outputs accompany this "
    "document for full reproducibility.", size=10)

doc.save("MTVS_Results_Section.docx")
print("saved MTVS_Results_Section.docx")
